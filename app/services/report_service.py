import os
import uuid
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Optional

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Image
from reportlab.lib.enums import TA_CENTER

from app.models.risk_profile import RiskAssessment, CustomRiskAssessment, RiskQuestionnaire
from app.models.client import ClientProfile
from app.models.ia_master import IAMaster
from sqlalchemy.orm import Session
from sqlalchemy import select
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from app.services.questionnaire_constants import QUESTIONNAIRE_DATA
from app.services.risk_profile_service import SCORING_RULES
from app.utils.encryption import decrypt_string

def resolve_logo_path(logo_path: Optional[str]) -> Optional[str]:
    """Try multiple strategies to find the logo file on disk, matching financial_report_generator.py logic."""
    if not logo_path:
        return None
        
    # Strategy 1: Absolute path
    if os.path.isabs(logo_path) and os.path.exists(logo_path):
        return logo_path
        
    # Strategy 2: Relative to CWD
    if os.path.exists(logo_path):
        return os.path.abspath(logo_path)

    # Strategy 3: Relative to backend root
    file_dir = os.path.dirname(os.path.abspath(__file__))
    # report_service.py is in app/services/, so backend_root is 2 levels up
    backend_root = os.path.abspath(os.path.join(file_dir, '..', '..'))
    joined_path = os.path.join(backend_root, logo_path)
    if os.path.exists(joined_path):
        return joined_path

    # Strategy 4: Try prepending 'uploads/'
    if not logo_path.startswith('uploads/'):
        uploads_path = os.path.join(backend_root, 'uploads', logo_path)
        if os.path.exists(uploads_path):
            return uploads_path

    return None

class ReportService:
    @staticmethod
    def _draw_footer(canvas, doc):
        canvas.saveState()
        # Header text
        canvas.setFont('Helvetica-Bold', 8)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(letter[0] - 0.5*inch, letter[1] - 0.5*inch, "STRICTLY CONFIDENTIAL")
        
        # Footer text
        canvas.setFont('Helvetica', 8)
        canvas.drawCentredString(letter[0]/2, 0.5*inch, f"Page {doc.page}")
        canvas.restoreState()

    @staticmethod
    def _add_page_number(run):
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(ns.qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    @staticmethod
    def generate_risk_profile_pdf(db: Session, assessment_id: uuid.UUID, ia_logo_override: str = None) -> BytesIO:
        # 1. Fetch Assessment and related data
        assessment = db.execute(
            select(RiskAssessment).where(RiskAssessment.id == assessment_id)
        ).scalar_one_or_none()
        
        if not assessment:
            raise ValueError("Assessment not found")
            
        client = assessment.client
        
        # Fetch IA details based on advisor name or reg number
        # Note: In the standalone, it used ia_registration_number. 
        # In current models, we look up IAMaster.
        # Fetch IA details based on advisor registration number
        ia = db.execute(
            select(IAMaster).where(IAMaster.ia_registration_number == client.advisor_registration_number)
        ).scalar_one_or_none()
        
        # Fallback if IA not found by registration number
        if not ia:
            ia = db.first(select(IAMaster))

        # Decrypt Entity Name
        entity_name = "N/A"
        if ia:
            raw_entity_name = ia.name_of_entity or ia.name_of_ia
            entity_name = decrypt_string(raw_entity_name) if raw_entity_name else "N/A"

        buffer = BytesIO()
        
        # 2. Setup Document
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=1.2*inch,
            bottomMargin=0.7*inch, 
            leftMargin=0.5*inch, 
            rightMargin=0.5*inch
        )
        
        story = []
        styles = getSampleStyleSheet()

        # Custom styles (Ported from risk_profile.py)
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e3c72'),
            spaceAfter=20,
            alignment=1,
            fontName='Helvetica-Bold'
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2a5298'),
            spaceAfter=12,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )

        normal_style = ParagraphStyle(
            'NormalCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica'
        )

        bold_style = ParagraphStyle(
            'BoldCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica-Bold'
        )

        preformatted_style = ParagraphStyle(
            'PreformattedStyle',
            parent=styles['Normal'],
            fontSize=8,
            fontName='Courier',
            leading=10,
            spaceAfter=2,
            leftIndent=0,
            rightIndent=0,
            alignment=0
        )

        # 3. Create Cover Page
        cover_title_style = ParagraphStyle(
            'CoverTitle',
            parent=title_style,
            fontSize=28,
            alignment=TA_CENTER,
            spaceBefore=100
        )
        cover_subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=normal_style,
            fontSize=14,
            alignment=TA_CENTER,
            spaceBefore=30,
            textColor=colors.grey
        )
        
        # Logo handling for cover
        raw_logo_path = ia_logo_override or (ia.ia_logo_path if ia else None)
        decrypted_logo_path = decrypt_string(raw_logo_path) if raw_logo_path else None
        resolved_logo = resolve_logo_path(decrypted_logo_path)
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=2.5*inch, height=1.25*inch, kind='proportional')
                logo.hAlign = 'CENTER'
                story.append(Spacer(1, 50))
                story.append(logo)
                story.append(Spacer(1, 10))
            except Exception as e:
                print(f"Error rendering logo in risk PDF: {e}")
        else:
            story.append(Spacer(1, 100))
        
        story.append(Paragraph("RISK PROFILE ASSESSMENT REPORT", cover_title_style))
        story.append(Spacer(1, 50))
        
        story.append(Paragraph(f"<b>CLIENT NAME:</b> {client.client_name}", cover_subtitle_style))
        story.append(Paragraph(f"<b>CLIENT CODE:</b> {client.client_code}", cover_subtitle_style))
        story.append(Paragraph(f"<b>ENTITY:</b> {entity_name}", cover_subtitle_style))
        story.append(Paragraph(f"<b>DATE:</b> {assessment.assessment_timestamp.strftime('%B %d, %Y')}", cover_subtitle_style))
        
        story.append(PageBreak())

        # 4. Result Summary Section
        story.append(Paragraph("ASSESSMENT SUMMARY", heading_style))
        summary_data = [
            ["Client Name:", client.client_name],
            ["Client Code:", client.client_code],
            ["Date of Assessment:", assessment.assessment_timestamp.strftime("%Y-%m-%d %H:%M")],
            ["Total Score:", str(assessment.calculated_score)],
            ["Risk Category:", assessment.assigned_risk_tier]
        ]
        summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
        summary_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
            ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))

        # 5. Questionnaire Responses
        story.append(Paragraph("QUESTIONNAIRE RESPONSES", heading_style))
        q_scores = assessment.question_scores or {}
        
        for q_id, q_data in QUESTIONNAIRE_DATA.items():
            s_info = q_scores.get(q_id, {})
            score_text = f" (Score: {s_info.get('score', 0)}/{s_info.get('max', 0)})"
            story.append(Paragraph(f"<b>{q_id.upper()}. {q_data['title']}{score_text}</b>", normal_style))
            story.append(Paragraph(q_data['question'], normal_style))
            
            # Handle Q2 (Factors) specifically
            if q_id == 'q2':
                q2_answers = assessment.q2_importance_factors
                q2_scores = s_info.get('details', {})
                table_data = [["Factor", "Importance", "Score"]]
                for f_code, f_name in q_data['factors'].items():
                    ans_code = q2_answers.get(f_code, "-")
                    ans_text = q_data['options'].get(ans_code, ans_code)
                    f_score = q2_scores.get(f_code, {}).get('score', 0)
                    table_data.append([f_name, Paragraph(f"<b>{ans_text}</b>", bold_style), str(f_score)])
                
                q2_table = Table(table_data, colWidths=[3*inch, 2*inch, 1*inch])
                q2_table.setStyle(TableStyle([
                    ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                    ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
                    ('PADDING', (0,0), (-1,-1), 4),
                ]))
                story.append(q2_table)
            else:
                # Standard questions - mapping q_id to model field name
                field_map = {
                    'q1': 'q1_interest_choice',
                    'q3': 'q3_probability_bet',
                    'q4': 'q4_portfolio_choice',
                    'q5': 'q5_loss_behavior',
                    'q6': 'q6_market_reaction',
                    'q7': 'q7_fund_selection',
                    'q8': 'q8_experience_level',
                    'q9': 'q9_time_horizon',
                    'q10': 'q10_net_worth',
                    'q11': 'q11_age_range',
                    'q12': 'q12_income_range',
                    'q13': 'q13_expense_range',
                    'q14': 'q14_dependents',
                    'q15': 'q15_active_loan',
                    'q16': 'q16_investment_objective'
                }
                field_name = field_map.get(q_id, q_id)
                selected_val = getattr(assessment, field_name, "N/A")
                
                options_str = ""
                for opt_code, opt_text in q_data['options'].items():
                    is_selected = opt_code.lower() == str(selected_val).lower()
                    prefix = f"<b>[✓]</b> " if is_selected else "[  ] "
                    text_color = "#006400" if is_selected else "#000000" # Dark Green for selected
                    options_str += f"<font color='{text_color}'>{prefix}{opt_code}. {opt_text}</font><br/>"
                
                story.append(Paragraph(options_str, ParagraphStyle('Options', parent=normal_style, leftIndent=20, leading=12)))
            
            story.append(Spacer(1, 10))
        
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"<b>TOTAL ASSESSMENT SCORE: {assessment.calculated_score} / 100</b>", bold_style))
        story.append(Spacer(1, 15))

        # 5a. Scoring Reference Chart
        story.append(Paragraph("SCORING REFERENCE", heading_style))
        ref_data = [["Q #", "Option", "Points"]]
        for q_id, rules in SCORING_RULES.items():
            if q_id == 'q2': continue # Handle Q2 separately
            for opt, score in rules.items():
                ref_data.append([q_id.upper(), f"Option {opt.upper()}", str(score)])
        
        # Split into multiple tables or one long one? One long one with columns
        # To make it compact, let's group by question
        compact_ref = [["Question", "Scoring Rule (Option: Points)"]]
        for q_id, rules in SCORING_RULES.items():
            if q_id == 'q2': continue
            rule_str = ", ".join([f"{opt.upper()}: {score}" for opt, score in rules.items()])
            compact_ref.append([q_id.upper(), rule_str])
        
        ref_table = Table(compact_ref, colWidths=[1.5*inch, 4.5*inch])
        ref_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(ref_table)
        
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Q2 Scoring Factors:</b> (A: Very, B: Somewhat, C: Not Important)", normal_style))
        q2_rules_data = [["Factor", "A", "B", "C"]]
        from app.services.questionnaire_constants import QUESTIONNAIRE_DATA as Q_DATA
        for f_id, f_name in Q_DATA['q2']['factors'].items():
            f_rules = SCORING_RULES['q2'].get(f_id, {})
            q2_rules_data.append([f_name, str(f_rules.get('A', 0)), str(f_rules.get('B', 0)), str(f_rules.get('C', 0))])
            
        q2_ref_table = Table(q2_rules_data, colWidths=[3*inch, 1*inch, 1*inch, 1*inch])
        q2_ref_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 7),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
        ]))
        story.append(q2_ref_table)
        story.append(Spacer(1, 15))
        
        # 6. Recommendation and Notes
        story.append(Paragraph("ADVISOR RECOMMENDATION", heading_style))
        story.append(Paragraph(f"<b>Classification:</b> {assessment.assigned_risk_tier}", normal_style))
        story.append(Paragraph(assessment.tier_recommendation or "No recommendation provided.", normal_style))
        
        # Add free space for manual recommendation
        story.append(Spacer(1, 12))
        story.append(Paragraph("Additional Advisor Guidance:", bold_style))
        story.append(Spacer(1, 120)) # Significantly larger space for manual writing
        # story.append(Paragraph("__________________________________________________________________________________________", normal_style))
        story.append(Spacer(1, 20))
        
        # DISCUSSION NOTES ALWAYS ON NEW PAGE AS REQUESTED
        story.append(PageBreak())

        if assessment.discussion_notes:
            story.append(Paragraph("DISCUSSION NOTES", heading_style))
            # Convert newlines for ReportLab Paragraph
            notes_html = assessment.discussion_notes.replace('\n', '<br/>')
            story.append(Paragraph(notes_html, normal_style))

        if assessment.disclaimer_text:
            story.append(Spacer(1, 20))
            story.append(Paragraph("DISCLAIMER", bold_style))
            # Convert newlines for ReportLab Paragraph
            discl_html = assessment.disclaimer_text.replace('\n', '<br/>')
            story.append(Paragraph(discl_html, ParagraphStyle('Disc', parent=normal_style, fontSize=7, textColor=colors.grey)))

        # 7. Signature Section
        story.append(Spacer(1, 40))
        sig_date = assessment.assessment_timestamp.strftime("%Y-%m-%d")
        
        sig_data = [
            [Paragraph("<b>__________________________</b><br/>Client Signature", normal_style), 
             Paragraph("<b>__________________________</b><br/>IA Advisor Signature", normal_style)],
            [Paragraph(f"Date: {sig_date}", normal_style),
             Paragraph(f"Date: {sig_date}", normal_style)]
        ]
        sig_table = Table(sig_data, colWidths=[3*inch, 3*inch])
        sig_table.setStyle(TableStyle([
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ]))
        story.append(sig_table)

        # 8. Build PDF
        doc.build(story, onFirstPage=ReportService._draw_footer, onLaterPages=ReportService._draw_footer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_risk_profile_docx(db: Session, assessment_id: uuid.UUID, ia_logo_override: str = None) -> BytesIO:
        # 1. Fetch Data
        assessment = db.execute(
            select(RiskAssessment).where(RiskAssessment.id == assessment_id)
        ).scalar_one_or_none()
        
        if not assessment:
            raise ValueError("Assessment not found")
            
        client = assessment.client
        ia = db.execute(
            select(IAMaster).where(IAMaster.ia_registration_number == client.advisor_registration_number)
        ).scalar_one_or_none()
        if not ia:
            ia = db.first(select(IAMaster))

        # Decrypt Entity Name
        entity_name = "N/A"
        if ia:
            raw_entity_name = ia.name_of_entity or ia.name_of_ia
            entity_name = decrypt_string(raw_entity_name) if raw_entity_name else "N/A"

        # 2. Setup DOCX
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        # 3. Header and Footer (Strictly Confidential & Page Numbers)
        section = doc.sections[0]
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "STRICTLY CONFIDENTIAL"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.style.font.bold = True
        header_p.style.font.size = Pt(8)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Page ")
        ReportService._add_page_number(footer_p.add_run())

        # 3a. Cover Page logic
        raw_logo_path = ia_logo_override or (ia.ia_logo_path if ia else None)
        decrypted_logo_path = decrypt_string(raw_logo_path) if raw_logo_path else None
        resolved_logo = resolve_logo_path(decrypted_logo_path)
        if resolved_logo:
            doc.add_picture(resolved_logo, width=Inches(2.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Separator for cover titles
        for _ in range(5): doc.add_paragraph()
        
        title = doc.add_heading('RISK PROFILE ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"\n\n\nCLIENT NAME: {client.client_name}\n")
        run.bold = True
        run.font.size = Pt(14)
        run = subtitle.add_run(f"CLIENT CODE: {client.client_code}\n")
        run.bold = True
        run = subtitle.add_run(f"ENTITY: {entity_name}\n")
        run = subtitle.add_run(f"DATE: {assessment.assessment_timestamp.strftime('%B %d, %Y')}")
        
        doc.add_page_break()

        # 4. Header
        header_title = doc.add_heading('RISK PROFILING REPORT', 0)
        header_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4. Summary Table
        doc.add_heading('ASSESSMENT SUMMARY', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        summary_rows = [
            ("Client Name:", client.client_name),
            ("Client Code:", client.client_code),
            ("Date of Assessment:", assessment.assessment_timestamp.strftime("%Y-%m-%d %H:%M")),
            ("Total Score:", str(assessment.calculated_score)),
            ("Risk Category:", assessment.assigned_risk_tier)
        ]
        
        for i, (label, value) in enumerate(summary_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_paragraph().add_run().add_break()

        # 5. Questionnaire Responses
        doc.add_heading('QUESTIONNAIRE RESPONSES', level=1)
        q_scores = assessment.question_scores or {}
        
        for q_id, q_data in QUESTIONNAIRE_DATA.items():
            s_info = q_scores.get(q_id, {})
            score_text = f" (Score: {s_info.get('score', 0)}/{s_info.get('max', 0)})"
            p = doc.add_paragraph()
            p.add_run(f"{q_id.upper()}. {q_data['title']}{score_text}").bold = True
            doc.add_paragraph(q_data['question'])
            
            if q_id == 'q2':
                q2_answers = assessment.q2_importance_factors
                q2_scores = s_info.get('details', {})
                t2 = doc.add_table(rows=1, cols=3)
                t2.style = 'Table Grid'
                hdr_cells = t2.rows[0].cells
                hdr_cells[0].text = 'Factor'
                hdr_cells[1].text = 'Importance'
                hdr_cells[2].text = 'Score'
                
                for f_code, f_name in q_data['factors'].items():
                    ans_code = q2_answers.get(f_code, "-")
                    ans_text = q_data['options'].get(ans_code, ans_code)
                    f_score = q2_scores.get(f_code, {}).get('score', 0)
                    row_cells = t2.add_row().cells
                    row_cells[0].text = f_name
                    row_cells[1].text = ans_text
                    row_cells[1].paragraphs[0].runs[0].bold = True
                    row_cells[2].text = str(f_score)
            else:
                # Standard questions - mapping q_id to model field name
                field_map = {
                    'q1': 'q1_interest_choice',
                    'q3': 'q3_probability_bet',
                    'q4': 'q4_portfolio_choice',
                    'q5': 'q5_loss_behavior',
                    'q6': 'q6_market_reaction',
                    'q7': 'q7_fund_selection',
                    'q8': 'q8_experience_level',
                    'q9': 'q9_time_horizon',
                    'q10': 'q10_net_worth',
                    'q11': 'q11_age_range',
                    'q12': 'q12_income_range',
                    'q13': 'q13_expense_range',
                    'q14': 'q14_dependents',
                    'q15': 'q15_active_loan',
                    'q16': 'q16_investment_objective'
                }
                field_name = field_map.get(q_id, q_id)
                selected_val = getattr(assessment, field_name, "N/A")

                for opt_code, opt_text in q_data['options'].items():
                    is_selected = opt_code.lower() == str(selected_val).lower()
                    prefix = f"[✓] " if is_selected else "[  ] "
                    option_p = doc.add_paragraph()
                    option_p.paragraph_format.left_indent = Pt(20)
                    run = option_p.add_run(f"{prefix}{opt_code}. {opt_text}")
                    if is_selected:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green
            
            doc.add_paragraph() # Spacer
        
        # Total Score Summary
        total_p = doc.add_paragraph()
        total_run = total_p.add_run(f"TOTAL ASSESSMENT SCORE: {assessment.calculated_score} / 100")
        total_run.bold = True
        total_run.font.size = Pt(11)
        doc.add_paragraph() # Extra spacer

        # 5a. Scoring Reference Chart
        doc.add_heading('SCORING REFERENCE', level=1)
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.style = 'Table Grid'
        hdr_cells = ref_table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Scoring Rule (Option: Points)'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].paragraphs[0].runs[0].bold = True

        for q_id, rules in SCORING_RULES.items():
            if q_id == 'q2': continue
            row_cells = ref_table.add_row().cells
            row_cells[0].text = q_id.upper()
            row_cells[1].text = ", ".join([f"{opt.upper()}: {score}" for opt, score in rules.items()])
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Q2 Scoring Factors: (A: Very, B: Somewhat, C: Not Important)").bold = True
        p.style.font.size = Pt(9)
        
        from app.services.questionnaire_constants import QUESTIONNAIRE_DATA as Q_DATA
        q2_table = doc.add_table(rows=1, cols=4)
        q2_table.style = 'Table Grid'
        hdr_cells = q2_table.rows[0].cells
        hdr_cells[0].text = 'Factor'
        hdr_cells[1].text = 'A'
        hdr_cells[2].text = 'B'
        hdr_cells[3].text = 'C'
        for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

        for f_id, f_name in Q_DATA['q2']['factors'].items():
            f_rules = SCORING_RULES['q2'].get(f_id, {})
            row_cells = q2_table.add_row().cells
            row_cells[0].text = f_name
            row_cells[1].text = str(f_rules.get('A', 0))
            row_cells[2].text = str(f_rules.get('B', 0))
            row_cells[3].text = str(f_rules.get('C', 0))
            for cell in row_cells: cell.paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph() # Extra spacer

        # 6. Recommendation and Notes
        doc.add_heading('ADVISOR RECOMMENDATION', level=1)
        p = doc.add_paragraph()
        p.add_run("Classification: ").bold = True
        p.add_run(assessment.assigned_risk_tier)
        
        doc.add_paragraph(assessment.tier_recommendation or "No recommendation provided.")
        
        # Add free space for manual recommendation
        doc.add_paragraph().add_run("Additional Advisor Guidance:").bold = True
        for _ in range(8): doc.add_paragraph() # Significant blank space
        # doc.add_paragraph("________________________________________________________________________________")
        
        # DISCUSSION NOTES ALWAYS ON NEW PAGE AS REQUESTED
        # doc.add_page_break()

        if assessment.discussion_notes:
            doc.add_heading('DISCUSSION NOTES', level=1)
            p = doc.add_paragraph()
            lines = assessment.discussion_notes.split('\n')
            for i, line in enumerate(lines):
                p.add_run(line)
                if i < len(lines) - 1:
                    p.add_run().add_break()

        if assessment.disclaimer_text:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("DISCLAIMER: ")
            run.bold = True
            lines = assessment.disclaimer_text.split('\n')
            for i, line in enumerate(lines):
                p.add_run(line)
                if i < len(lines) - 1:
                    p.add_run().add_break()

        # 7. Signature Section
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.width = Inches(6)
        
        cells = sig_table.rows[0].cells
        cells[0].text = "__________________________\nClient Signature"
        cells[1].text = "__________________________\nIA Advisor Signature"
        
        sig_date = assessment.assessment_timestamp.strftime("%Y-%m-%d")
        cells = sig_table.rows[1].cells
        cells[0].text = f"Date: {sig_date}"
        cells[1].text = f"Date: {sig_date}"

        # 8. Save to Buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_custom_risk_profile_pdf(db: Session, assessment_id: uuid.UUID, ia_logo_override: str = None) -> BytesIO:
        def fmt_score(val):
            try:
                f_val = float(val)
                return int(f_val) if f_val % 1 == 0 else f_val
            except: return val

        # 1. Fetch Data
        assessment = db.execute(
            select(CustomRiskAssessment).where(CustomRiskAssessment.id == assessment_id)
        ).scalar_one_or_none()
        
        if not assessment:
            raise ValueError("Custom risk assessment not found")
        
        questionnaire = assessment.questionnaire
        client = assessment.client
        
        ia = db.execute(
            select(IAMaster).where(IAMaster.ia_registration_number == client.advisor_registration_number)
        ).scalar_one_or_none()
        if not ia:
            ia = db.first(select(IAMaster))

        # Decrypt Entity Name
        entity_name = "N/A"
        if ia:
            raw_entity_name = ia.name_of_entity or ia.name_of_ia
            entity_name = decrypt_string(raw_entity_name) if raw_entity_name else "N/A"

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=1.2*inch,
            bottomMargin=0.7*inch, 
            leftMargin=0.5*inch, 
            rightMargin=0.5*inch
        )
        
        story = []
        styles = getSampleStyleSheet()

        # Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e3c72'),
            spaceAfter=20,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2a5298'),
            spaceAfter=12,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'NormalCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica'
        )
        bold_style = ParagraphStyle(
            'BoldCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica-Bold'
        )

        cover_title_style = ParagraphStyle(
            'CoverTitle',
            parent=title_style,
            fontSize=28,
            alignment=TA_CENTER,
            spaceBefore=100
        )
        cover_subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=normal_style,
            fontSize=14,
            alignment=TA_CENTER,
            spaceBefore=30,
            textColor=colors.grey
        )

        # 3. Create Cover Page
        raw_logo_path = ia_logo_override or (ia.ia_logo_path if ia else None)
        decrypted_logo_path = decrypt_string(raw_logo_path) if raw_logo_path else None
        resolved_logo = resolve_logo_path(decrypted_logo_path)
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=2.5*inch, height=1.25*inch, kind='proportional')
                logo.hAlign = 'CENTER'
                story.append(Spacer(1, 50))
                story.append(logo)
                story.append(Spacer(1, 10))
            except:
                story.append(Spacer(1, 100))
        else:
            story.append(Spacer(1, 100))
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("RISK PROFILE ASSESSMENT REPORT", cover_title_style))
        story.append(Spacer(1, 50))
        
        story.append(Paragraph(f"<b>CLIENT NAME:</b> {client.client_name}", cover_subtitle_style))
        story.append(Paragraph(f"<b>CLIENT CODE:</b> {client.client_code}", cover_subtitle_style))
        story.append(Paragraph(f"<b>ENTITY:</b> {entity_name}", cover_subtitle_style))
        story.append(Paragraph(f"<b>DATE:</b> {assessment.submitted_at.strftime('%B %d, %Y')}", cover_subtitle_style))
        
        story.append(PageBreak())

        # 4. Result Summary Section
        story.append(Paragraph("ASSESSMENT SUMMARY", heading_style))
        summary_data = [
            ["Client Name:", client.client_name],
            ["Client Code:", client.client_code],
            ["Date of Assessment:", assessment.submitted_at.strftime("%Y-%m-%d %H:%M")],
            ["Total Score:", f"{fmt_score(assessment.total_score)} / {fmt_score(questionnaire.max_possible_score)}"],
            ["Risk Category:", assessment.category_name or "N/A"]
        ]
        summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
        summary_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
            ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))

        # 5. Questionnaire Responses
        story.append(Paragraph("QUESTIONNAIRE RESPONSES", heading_style))
        
        responses = assessment.responses or {}
        for q_idx, q_data in enumerate(questionnaire.questions):
            q_id = q_data.get('id', f'q_{q_idx}')
            ans_info = responses.get(q_id, {})
            score_text = f" (Score: {fmt_score(ans_info.get('score', 0))})"
            story.append(Paragraph(f"<b>{q_idx+1}. {q_data.get('text', '')}{score_text}</b>", normal_style))
            
            selected_option_id = ans_info.get('option_id')
            options_p = ""
            for opt in q_data.get('options', []):
                opt_id = opt.get('id')
                is_selected = str(opt_id) == str(selected_option_id)
                prefix = f"<b>[✓]</b> " if is_selected else "[  ] "
                text_color = "#006400" if is_selected else "#000000"
                options_p += f"<font color='{text_color}'>{prefix}{opt.get('text', '')}</font><br/>"
            
            story.append(Paragraph(options_p, ParagraphStyle('Options', parent=normal_style, leftIndent=20, leading=12)))
            story.append(Spacer(1, 10))

        # 5a. Scoring Reference Chart
        story.append(Paragraph("SCORING REFERENCE", heading_style))
        cat_data = [["Category", "Score Range", "Description"]]
        for cat in (questionnaire.categories or []):
            cat_data.append([
                cat.get('name', 'N/A'),
                f"{fmt_score(cat.get('min_score'))} - {fmt_score(cat.get('max_score'))}",
                Paragraph(cat.get('description', ''), normal_style)
            ])
        
        cat_table = Table(cat_data, colWidths=[1.5*inch, 1.5*inch, 3*inch])
        cat_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(cat_table)
        
        story.append(Spacer(1, 10))
        story.append(Paragraph("Question Scoring Rules", normal_style))
        q_ref_data = [["Question", "Options & Points"]]
        for q_idx, q in enumerate(questionnaire.questions):
            # Using letters (A, B, C...) for options to match standard report style
            rule_str = ", ".join([f"{chr(65 + i)}: {fmt_score(o.get('score'))}" for i, o in enumerate(q.get('options', []))])
            q_ref_data.append([f"Q{q_idx+1}", Paragraph(rule_str, normal_style)])
        
        q_ref_table = Table(q_ref_data, colWidths=[1*inch, 5*inch])
        q_ref_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 7),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(q_ref_table)

        # 6. Recommendation and Notes
        story.append(Spacer(1, 15))
        story.append(Paragraph("ADVISOR GUIDANCE", heading_style))
        story.append(Paragraph(f"<b>Classification:</b> {assessment.category_name or 'N/A'}", normal_style))
        
        if assessment.discussion_notes:
            story.append(Spacer(1, 10))
            story.append(Paragraph("DISCUSSION NOTES", bold_style))
            notes_html = assessment.discussion_notes.replace('\n', '<br/>')
            story.append(Paragraph(notes_html, normal_style))

        # 7. Signature Section
        story.append(Spacer(1, 40))
        sig_date = assessment.submitted_at.strftime("%Y-%m-%d")
        sig_data = [
            [Paragraph("<b>__________________________</b><br/>Client Signature", normal_style), 
             Paragraph("<b>__________________________</b><br/>IA Advisor Signature", normal_style)],
            [Paragraph(f"Date: {sig_date}", normal_style),
             Paragraph(f"Date: {sig_date}", normal_style)]
        ]
        sig_table = Table(sig_data, colWidths=[3*inch, 3*inch])
        sig_table.setStyle(TableStyle([
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ]))
        story.append(sig_table)

        doc.build(story, onFirstPage=ReportService._draw_footer, onLaterPages=ReportService._draw_footer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_custom_risk_profile_docx(db: Session, assessment_id: uuid.UUID, ia_logo_override: str = None) -> BytesIO:
        def fmt_score(val):
            try:
                f_val = float(val)
                return int(f_val) if f_val % 1 == 0 else f_val
            except: return val

        # 1. Fetch Data
        assessment = db.execute(
            select(CustomRiskAssessment).where(CustomRiskAssessment.id == assessment_id)
        ).scalar_one_or_none()
        
        if not assessment:
            raise ValueError("Custom risk assessment not found")
        
        questionnaire = assessment.questionnaire
        client = assessment.client
        ia = db.execute(select(IAMaster)).first()
        if ia: ia = ia[0]

        # 2. Setup DOCX
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        # Header/Footer
        section = doc.sections[0]
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "STRICTLY CONFIDENTIAL"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.style.font.bold = True
        header_p.style.font.size = Pt(8)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Page ")
        ReportService._add_page_number(footer_p.add_run())

        # Cover Page
        resolved_logo = resolve_logo_path(ia_logo_override or (ia.ia_logo_path if ia else None))
        if resolved_logo:
            doc.add_picture(resolved_logo, width=Inches(2.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for _ in range(5): doc.add_paragraph()
        title = doc.add_heading('RISK PROFILE ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"\n\n\nCLIENT NAME: {client.client_name}\n")
        run.bold = True
        run.font.size = Pt(14)
        run = subtitle.add_run(f"CLIENT CODE: {client.client_code}\n")
        if ia:
            run = subtitle.add_run(f"ENTITY: {ia.name_of_entity or ia.name_of_ia}\n")
        run = subtitle.add_run(f"DATE: {assessment.submitted_at.strftime('%B %d, %Y')}")
        
        doc.add_page_break()

        # Summary
        doc.add_heading('ASSESSMENT SUMMARY', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        summary_rows = [
            ("Client Name:", client.client_name),
            ("Client Code:", client.client_code),
            ("Date of Assessment:", assessment.submitted_at.strftime("%Y-%m-%d %H:%M")),
            ("Total Score:", f"{fmt_score(assessment.total_score)} / {fmt_score(questionnaire.max_possible_score)}"),
            ("Risk Category:", assessment.category_name or "N/A")
        ]
        for i, (label, value) in enumerate(summary_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_paragraph().add_run().add_break()

        # Responses
        doc.add_heading('QUESTIONNAIRE RESPONSES', level=1)
        responses = assessment.responses or {}
        for q_idx, q_data in enumerate(questionnaire.questions):
            q_id = q_data.get('id', f'q_{q_idx}')
            ans_info = responses.get(q_id, {})
            score_text = f" (Score: {fmt_score(ans_info.get('score', 0))})"
            
            p = doc.add_paragraph()
            p.add_run(f"{q_idx+1}. {q_data.get('text', '')}{score_text}").bold = True
            
            selected_option_id = ans_info.get('option_id')
            for opt in q_data.get('options', []):
                opt_id = opt.get('id')
                is_selected = str(opt_id) == str(selected_option_id)
                prefix = f"[✓] " if is_selected else "[  ] "
                
                option_p = doc.add_paragraph()
                option_p.paragraph_format.left_indent = Pt(20)
                run = option_p.add_run(f"{prefix}{opt.get('text', '')}")
                if is_selected:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 0)
            
            doc.add_paragraph()

        # Score Reference Table
        doc.add_heading('SCORING REFERENCE', level=1)
        cat_table = doc.add_table(rows=1, cols=3)
        cat_table.style = 'Table Grid'
        hdr_cells = cat_table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Score Range'
        hdr_cells[2].text = 'Description'
        for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

        for cat in (questionnaire.categories or []):
            row_cells = cat_table.add_row().cells
            row_cells[0].text = cat.get('name', 'N/A')
            row_cells[1].text = f"{fmt_score(cat.get('min_score'))} - {fmt_score(cat.get('max_score'))}"
            row_cells[2].text = cat.get('description', '')

        doc.add_paragraph().add_run("Question Scoring Rules").bold = True
        q_ref_table = doc.add_table(rows=1, cols=2)
        q_ref_table.style = 'Table Grid'
        q_ref_table.rows[0].cells[0].text = 'Question'
        q_ref_table.rows[0].cells[1].text = 'Options & Points'
        for cell in q_ref_table.rows[0].cells: cell.paragraphs[0].runs[0].bold = True

        for q_idx, q in enumerate(questionnaire.questions):
            row_cells = q_ref_table.add_row().cells
            row_cells[0].text = f"Q{q_idx+1}"
            row_cells[1].text = ", ".join([f"{chr(65 + i)}: {fmt_score(o.get('score'))}" for i, o in enumerate(q.get('options', []))])
            for cell in row_cells: cell.paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph()

        # Guidance
        doc.add_heading('ADVISOR GUIDANCE', level=1)
        doc.add_paragraph().add_run(f"Classification: {assessment.category_name or 'N/A'}").bold = True
        
        if assessment.discussion_notes:
            doc.add_paragraph().add_run("DISCUSSION NOTES:").bold = True
            doc.add_paragraph(assessment.discussion_notes)

        # Signatures
        doc.add_paragraph().add_run().add_break()
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.width = Inches(6)
        sig_table.rows[0].cells[0].text = "__________________________\nClient Signature"
        sig_table.rows[0].cells[1].text = "__________________________\nIA Advisor Signature"
        sig_date = assessment.submitted_at.strftime("%Y-%m-%d")
        sig_table.rows[1].cells[0].text = f"Date: {sig_date}"
    @staticmethod
    def generate_blank_risk_form_pdf(
        db: Optional[Session], 
        questionnaire_id: str, 
        ia_logo_override: str = None, 
        ia_data: dict = None,
        questionnaire_data: dict = None
    ) -> BytesIO:
        # Fetch IA details 
        if ia_data:
            ia = type('MockIA', (), ia_data)() # Convert dict to object-like access if needed, or adapt code
            # Actually, let's just adapt the code below to handle both.
        else:
            ia = db.execute(select(IAMaster)).scalar_one_or_none() if db else None

        if questionnaire_id == "sample-form":
            # Mock questionnaire object for legacy form
            from app.services.questionnaire_constants import QUESTIONNAIRE_DATA
            
            mock_questions = []
            for q_key, q_val in QUESTIONNAIRE_DATA.items():
                if q_key == 'q2':
                    # Main Question 2 Instruction
                    mock_questions.append({
                        'text': q_val.get('question', ''),
                        'options': [], # No options for the instruction itself
                        'is_main': True
                    })
                    # Sub-questions 2.1 to 2.8
                    options = [{'text': v} for v in q_val.get('options', {}).values()]
                    for i, (f_key, f_val) in enumerate(q_val.get('factors', {}).items()):
                        mock_questions.append({
                            'text': f_val,
                            'options': options,
                            'is_sub': True,
                            'sub_idx': i + 1
                        })
                else:
                    options = [{'text': v} for v in q_val.get('options', {}).values()]
                    mock_questions.append({
                        'text': q_val.get('question', ''),
                        'options': options,
                        'is_main': True
                    })

            class MockQuestionnaire:
                def __init__(self, questions):
                    self.id = "SYSTEM-DEFAULT"
                    self.questions = questions
                    self.disclaimer = None # Remove disclaimer for system default
            
            questionnaire = MockQuestionnaire(mock_questions)
        elif questionnaire_data:
            # Use data provided by caller (e.g. from Bridge)
            class BridgeQuestionnaire:
                def __init__(self, data):
                    self.id = data.get('id', 'N/A')
                    self.questions = data.get('questions') or []
                    # Bridge returns JSONB as list of dicts, make sure it is parsed
                    if isinstance(self.questions, str):
                        import json
                        self.questions = json.loads(self.questions)
                    self.disclaimer = data.get('disclaimer')
            
            questionnaire = BridgeQuestionnaire(questionnaire_data)
        else:
            # Fetch custom questionnaire from local DB
            try:
                q_uuid = uuid.UUID(questionnaire_id)
                questionnaire = db.execute(
                    select(RiskQuestionnaire).where(RiskQuestionnaire.id == q_uuid)
                ).scalar_one_or_none()
            except:
                raise ValueError("Questionnaire not found")
            
            if not questionnaire:
                raise ValueError("Questionnaire not found")

        # 1. Setup Document
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=1.2*inch,
            bottomMargin=0.7*inch, 
            leftMargin=0.5*inch, 
            rightMargin=0.5*inch
        )
        
        story = []
        styles = getSampleStyleSheet()

        # Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e3c72'),
            spaceAfter=20,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2a5298'),
            spaceAfter=12,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'NormalCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica'
        )
        bold_style = ParagraphStyle(
            'BoldCustom',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=4,
            fontName='Helvetica-Bold'
        )

        cover_title_style = ParagraphStyle(
            'CoverTitle',
            parent=title_style,
            fontSize=28,
            alignment=TA_CENTER,
            spaceBefore=100
        )
        cover_subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=normal_style,
            fontSize=14,
            alignment=TA_CENTER,
            spaceBefore=30,
            textColor=colors.grey
        )

        # 3. Create Cover Page
        ia_logo_path = ia_logo_override
        if not ia_logo_path:
            if ia_data:
                ia_logo_path = ia_data.get("ia_logo_path")
            elif ia:
                ia_logo_path = ia.ia_logo_path

        resolved_logo = resolve_logo_path(ia_logo_path)
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=2.5*inch, height=1.25*inch, kind='proportional')
                story.append(Spacer(1, 50))
                story.append(logo)
            except: pass
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("RISK ASSESSMENT FORM", cover_title_style))
        story.append(Paragraph(f"(FORM ID: {str(questionnaire.id)[:8].upper()})", cover_subtitle_style))
        story.append(Spacer(1, 50))
        
        story.append(Paragraph("<b>CLIENT NAME:</b> __________________________", cover_subtitle_style))
        story.append(Paragraph("<b>CLIENT CODE:</b> __________________________", cover_subtitle_style))
        
        # Get IA Identifiers
        if ia_data:
            entity_name = ia_data.get("name_of_entity") or ia_data.get("name_of_ia")
            reg_no = ia_data.get("ia_registration_number") or "N/A"
        elif ia:
            entity_name = ia.name_of_entity or ia.name_of_ia
            reg_no = ia.ia_registration_number or "N/A"
        else:
            entity_name = None
            reg_no = None

        if entity_name:
            story.append(Paragraph(f"<b>ENTITY:</b> {entity_name}", cover_subtitle_style))
        if reg_no:
            story.append(Paragraph(f"<b>REGISTRATION NO:</b> {reg_no}", cover_subtitle_style))
            
        story.append(Paragraph(f"<b>DATE:</b> __________________________", cover_subtitle_style))
        
        story.append(PageBreak())

        # 5. Questionnaire Items
        story.append(Paragraph("QUESTIONNAIRE", heading_style))
        
        main_q_idx = 0
        for q_data in questionnaire.questions:
            # Handle Numbering
            if q_data.get('is_sub'):
                label = f"{main_q_idx}.{q_data.get('sub_idx')}"
            else:
                main_q_idx += 1
                label = f"{main_q_idx}"
            
            story.append(Paragraph(f"<b>{label}. {q_data.get('text', '')}</b>", normal_style))
            
            options_p = ""
            for i, opt in enumerate(q_data.get('options', [])):
                prefix = f"[ {chr(65 + i)} ] "
                options_p += f"{prefix}{opt.get('text', '')}<br/>"
            
            if options_p:
                story.append(Paragraph(options_p, ParagraphStyle('Options', parent=normal_style, leftIndent=30, leading=12)))
            story.append(Spacer(1, 10))

        # 7. Disclaimer Section
        if questionnaire.disclaimer:
            story.append(Spacer(1, 20))
            story.append(Paragraph("DISCLAIMER", bold_style))
            story.append(Paragraph(questionnaire.disclaimer, normal_style))
            story.append(Spacer(1, 10))

        # 8. Discussion Notes Section
        story.append(Spacer(1, 20))
        story.append(Paragraph("DISCUSSION & STRATEGIC NOTES", bold_style))
        story.append(Spacer(1, 80)) # Reduced blank space for manual notes
        
        # 9. Signature Section
        story.append(Spacer(1, 30))
        sig_data = [
            [Paragraph("<b>__________________________</b><br/>Client Signature<br/><br/>Date: __________________", normal_style), 
             Paragraph("<b>__________________________</b><br/>IA Advisor Signature<br/><br/>Date: __________________", normal_style)]
        ]
        sig_table = Table(sig_data, colWidths=[3*inch, 3*inch])
        sig_table.setStyle(TableStyle([
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ]))
        story.append(sig_table)

        doc.build(story, onFirstPage=ReportService._draw_footer, onLaterPages=ReportService._draw_footer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_risk_profile_pdf_bridge(assessment_data: dict, client_data: dict, ia_data: dict, ia_logo_override: str = None) -> BytesIO:
        """
        Bridge-compatible version of generate_risk_profile_pdf.
        Accepts dictionaries instead of DB model instances.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=1.2*inch,
            bottomMargin=0.7*inch, 
            leftMargin=0.5*inch, 
            rightMargin=0.5*inch
        )
        
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#1e3c72'),
            spaceAfter=20, alignment=1, fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2a5298'),
            spaceAfter=12, spaceBefore=15, fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'NormalCustom', parent=styles['Normal'], fontSize=9, spaceAfter=4, fontName='Helvetica'
        )
        bold_style = ParagraphStyle(
            'BoldCustom', parent=styles['Normal'], fontSize=9, spaceAfter=4, fontName='Helvetica-Bold'
        )

        cover_title_style = ParagraphStyle('CoverTitle', parent=title_style, fontSize=28, alignment=TA_CENTER, spaceBefore=100)
        cover_subtitle_style = ParagraphStyle('CoverSubtitle', parent=normal_style, fontSize=14, alignment=TA_CENTER, spaceBefore=30, textColor=colors.grey)
        
        # Logo handling for cover
        resolved_logo = resolve_logo_path(ia_logo_override or ia_data.get('ia_logo_path'))
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=2.5*inch, height=1.25*inch, kind='proportional')
                story.append(Spacer(1, 50))
                story.append(logo)
                story.append(Spacer(1, 30))
            except: pass
        
        story.append(Paragraph("RISK PROFILE ASSESSMENT REPORT", cover_title_style))
        story.append(Spacer(1, 50))
        
        story.append(Paragraph(f"<b>CLIENT NAME:</b> {client_data.get('client_name', 'N/A')}", cover_subtitle_style))
        story.append(Paragraph(f"<b>CLIENT CODE:</b> {client_data.get('client_code', 'N/A')}", cover_subtitle_style))
        story.append(Paragraph(f"<b>ENTITY:</b> {ia_data.get('name_of_entity') or ia_data.get('name_of_ia', 'N/A')}", cover_subtitle_style))
        
        # Handle timestamp from bridge (string or datetime)
        ts = assessment_data.get('assessment_timestamp')
        date_str = "N/A"
        if ts:
            if isinstance(ts, str):
                try: 
                    # Try to parse ISO format
                    from dateutil import parser
                    ts_dt = parser.parse(ts)
                    date_str = ts_dt.strftime('%B %d, %Y')
                except: 
                    date_str = str(ts)
            elif hasattr(ts, 'strftime'):
                date_str = ts.strftime('%B %d, %Y')
        
        story.append(Paragraph(f"<b>DATE:</b> {date_str}", cover_subtitle_style))
        story.append(PageBreak())

        # Summary
        story.append(Paragraph("ASSESSMENT SUMMARY", heading_style))
        summary_data = [
            ["Client Name:", client_data.get('client_name')],
            ["Client Code:", client_data.get('client_code')],
            ["Date of Assessment:", date_str],
            ["Total Score:", str(assessment_data.get('calculated_score', 0))],
            ["Risk Category:", str(assessment_data.get('assigned_risk_tier', 'N/A'))]
        ]
        summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
        summary_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
            ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))

        # Questionnaire Responses
        story.append(Paragraph("QUESTIONNAIRE RESPONSES", heading_style))
        q_scores = assessment_data.get('question_scores') or {}
        
        from app.services.questionnaire_constants import QUESTIONNAIRE_DATA
        for q_id, q_data in QUESTIONNAIRE_DATA.items():
            s_info = q_scores.get(q_id, {})
            score_text = f" (Score: {s_info.get('score', 0)}/{s_info.get('max', 0)})"
            story.append(Paragraph(f"<b>{q_id.upper()}. {q_data['title']}{score_text}</b>", normal_style))
            story.append(Paragraph(q_data['question'], normal_style))
            
            if q_id == 'q2':
                q2_answers = assessment_data.get('q2_importance_factors') or {}
                q2_scores = s_info.get('details', {})
                table_data = [["Factor", "Importance", "Score"]]
                for f_code, f_name in q_data['factors'].items():
                    ans_code = q2_answers.get(f_code, "-")
                    ans_text = q_data['options'].get(ans_code, ans_code)
                    f_score = q2_scores.get(f_code, {}).get('score', 0)
                    table_data.append([f_name, Paragraph(f"<b>{ans_text}</b>", bold_style), str(f_score)])
                
                q2_table = Table(table_data, colWidths=[3*inch, 2*inch, 1*inch])
                q2_table.setStyle(TableStyle([
                    ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                    ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
                    ('PADDING', (0,0), (-1,-1), 4),
                ]))
                story.append(q2_table)
            else:
                field_map = {
                    'q1': 'q1_interest_choice', 'q3': 'q3_probability_bet', 'q4': 'q4_portfolio_choice',
                    'q5': 'q5_loss_behavior', 'q6': 'q6_market_reaction', 'q7': 'q7_fund_selection',
                    'q8': 'q8_experience_level', 'q9': 'q9_time_horizon', 'q10': 'q10_net_worth',
                    'q11': 'q11_age_range', 'q12': 'q12_income_range', 'q13': 'q13_expense_range',
                    'q14': 'q14_dependents', 'q15': 'q15_active_loan', 'q16': 'q16_investment_objective'
                }
                selected_val = assessment_data.get(field_map.get(q_id, q_id), "N/A")
                options_str = ""
                for opt_code, opt_text in q_data['options'].items():
                    is_selected = str(opt_code).lower() == str(selected_val).lower()
                    prefix = f"<b>[✓]</b> " if is_selected else "[  ] "
                    text_color = "#006400" if is_selected else "#000000"
                    options_str += f"<font color='{text_color}'>{prefix}{opt_code}. {opt_text}</font><br/>"
                story.append(Paragraph(options_str, ParagraphStyle('Options', parent=normal_style, leftIndent=20, leading=12)))
            story.append(Spacer(1, 10))
        
        story.append(Paragraph(f"<b>TOTAL ASSESSMENT SCORE: {assessment_data.get('calculated_score', 0)} / 100</b>", bold_style))
        story.append(Spacer(1, 15))

        # References
        story.append(Paragraph("SCORING REFERENCE", heading_style))
        compact_ref = [["Question", "Scoring Rule (Option: Points)"]]
        for q_id, rules in SCORING_RULES.items():
            if q_id == 'q2': continue
            rule_str = ", ".join([f"{opt.upper()}: {score}" for opt, score in rules.items()])
            compact_ref.append([q_id.upper(), rule_str])
        
        ref_table = Table(compact_ref, colWidths=[1.5*inch, 4.5*inch])
        ref_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(ref_table)

        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Q2 Scoring Factors:</b> (A: Very, B: Somewhat, C: Not Important)", normal_style))
        q2_rules_data = [["Factor", "A", "B", "C"]]
        from app.services.questionnaire_constants import QUESTIONNAIRE_DATA as Q_DATA
        for f_id, f_name in Q_DATA['q2']['factors'].items():
            f_rules = SCORING_RULES['q2'].get(f_id, {})
            q2_rules_data.append([f_name, str(f_rules.get('A', 0)), str(f_rules.get('B', 0)), str(f_rules.get('C', 0))])
            
        q2_ref_table = Table(q2_rules_data, colWidths=[3*inch, 1*inch, 1*inch, 1*inch])
        q2_ref_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,-1), 7),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
        ]))
        story.append(q2_ref_table)
        story.append(Spacer(1, 15))

        # Recommendations
        story.append(Paragraph("ADVISOR RECOMMENDATION", heading_style))
        story.append(Paragraph(f"<b>Classification:</b> {assessment_data.get('assigned_risk_tier', 'N/A')}", normal_style))
        story.append(Paragraph(assessment_data.get('tier_recommendation') or "No recommendation provided.", normal_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Additional Advisor Guidance:", bold_style))
        story.append(Spacer(1, 120))
        
        # Notes & Disclaimer
        story.append(PageBreak())
        if assessment_data.get('discussion_notes'):
            story.append(Paragraph("DISCUSSION NOTES", heading_style))
            story.append(Paragraph(assessment_data['discussion_notes'].replace('\n', '<br/>'), normal_style))

        if assessment_data.get('disclaimer_text'):
            story.append(Spacer(1, 20))
            story.append(Paragraph("DISCLAIMER", bold_style))
            story.append(Paragraph(assessment_data['disclaimer_text'].replace('\n', '<br/>'), ParagraphStyle('Disc', parent=normal_style, fontSize=7, textColor=colors.grey)))

        # Signatures
        story.append(Spacer(1, 40))
        sig_data = [
            [Paragraph("<b>__________________________</b><br/>Client Signature", normal_style), 
             Paragraph("<b>__________________________</b><br/>IA Advisor Signature", normal_style)],
            [Paragraph(f"Date: ________", normal_style),
             Paragraph(f"Date: ________", normal_style)]
        ]
        sig_table = Table(sig_data, colWidths=[3*inch, 3*inch])
        sig_table.setStyle(TableStyle([('TOPPADDING', (0,0), (-1,-1), 10), ('ALIGN', (0,0), (-1,-1), 'LEFT')]))
        story.append(sig_table)

        doc.build(story, onFirstPage=ReportService._draw_footer, onLaterPages=ReportService._draw_footer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_risk_profile_docx_bridge(assessment_data: dict, client_data: dict, ia_data: dict, ia_logo_override: str = None) -> BytesIO:
        """DOCX version for Bridge assessments, matching PDF layout/content."""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        # 3. Header and Footer (Strictly Confidential & Page Numbers)
        section = doc.sections[0]
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "STRICTLY CONFIDENTIAL"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.style.font.bold = True
        header_p.style.font.size = Pt(8)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Page ")
        ReportService._add_page_number(footer_p.add_run())

        # 3a. Cover Page logic
        resolved_logo = resolve_logo_path(ia_logo_override or ia_data.get('ia_logo_path'))
        if resolved_logo:
            try:
                doc.add_picture(resolved_logo, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass

        # Separator for cover titles
        for _ in range(5): doc.add_paragraph()
        
        title = doc.add_heading('RISK PROFILE ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"\n\n\nCLIENT NAME: {client_data.get('client_name')}\n")
        run.bold = True
        run.font.size = Pt(14)
        run = subtitle.add_run(f"CLIENT CODE: {client_data.get('client_code')}\n")
        run.bold = True
        run = subtitle.add_run(f"ENTITY: {ia_data.get('name_of_entity') or ia_data.get('name_of_ia')}\n")
        
        # Handle timestamp from bridge
        ts = assessment_data.get('assessment_timestamp')
        date_str = "N/A"
        if ts:
            if isinstance(ts, str):
                try: 
                    from dateutil import parser
                    ts_dt = parser.parse(ts)
                    date_str = ts_dt.strftime('%B %d, %Y')
                except: date_str = str(ts)
            elif hasattr(ts, 'strftime'):
                date_str = ts.strftime('%B %d, %Y')
        
        run = subtitle.add_run(f"DATE: {date_str}")
        
        doc.add_page_break()

        # 4. Summary Table
        doc.add_heading('ASSESSMENT SUMMARY', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        summary_rows = [
            ("Client Name:", client_data.get('client_name')),
            ("Client Code:", client_data.get('client_code')),
            ("Date of Assessment:", date_str),
            ("Total Score:", str(assessment_data.get('calculated_score', 0))),
            ("Risk Category:", str(assessment_data.get('assigned_risk_tier', 'N/A')))
        ]
        
        for i, (label, value) in enumerate(summary_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_paragraph().add_run().add_break()

        # 5. Questionnaire Responses
        doc.add_heading('QUESTIONNAIRE RESPONSES', level=1)
        q_scores = assessment_data.get('question_scores') or {}
        from app.services.questionnaire_constants import QUESTIONNAIRE_DATA
        
        for q_id, q_data in QUESTIONNAIRE_DATA.items():
            s_info = q_scores.get(q_id, {})
            score_text = f" (Score: {s_info.get('score', 0)}/{s_info.get('max', 0)})"
            p = doc.add_paragraph()
            p.add_run(f"{q_id.upper()}. {q_data['title']}{score_text}").bold = True
            doc.add_paragraph(q_data['question'])
            
            if q_id == 'q2':
                q2_answers = assessment_data.get('q2_importance_factors') or {}
                q2_scores = s_info.get('details', {})
                t2 = doc.add_table(rows=1, cols=3)
                t2.style = 'Table Grid'
                hdr_cells = t2.rows[0].cells
                hdr_cells[0].text = 'Factor'
                hdr_cells[1].text = 'Importance'
                hdr_cells[2].text = 'Score'
                
                for f_code, f_name in q_data['factors'].items():
                    ans_code = q2_answers.get(f_code, "-")
                    ans_text = q_data['options'].get(ans_code, ans_code)
                    f_score = q2_scores.get(f_code, {}).get('score', 0)
                    row_cells = t2.add_row().cells
                    row_cells[0].text = f_name
                    row_cells[1].text = ans_text
                    row_cells[1].paragraphs[0].runs[0].bold = True
                    row_cells[2].text = str(f_score)
            else:
                field_map = {
                    'q1': 'q1_interest_choice', 'q3': 'q3_probability_bet', 'q4': 'q4_portfolio_choice',
                    'q5': 'q5_loss_behavior', 'q6': 'q6_market_reaction', 'q7': 'q7_fund_selection',
                    'q8': 'q8_experience_level', 'q9': 'q9_time_horizon', 'q10': 'q10_net_worth',
                    'q11': 'q11_age_range', 'q12': 'q12_income_range', 'q13': 'q13_expense_range',
                    'q14': 'q14_dependents', 'q15': 'q15_active_loan', 'q16': 'q16_investment_objective'
                }
                selected_val = assessment_data.get(field_map.get(q_id, q_id), "N/A")

                for opt_code, opt_text in q_data['options'].items():
                    is_selected = str(opt_code).lower() == str(selected_val).lower()
                    prefix = f"[✓] " if is_selected else "[  ] "
                    option_p = doc.add_paragraph()
                    option_p.paragraph_format.left_indent = Pt(20)
                    run = option_p.add_run(f"{prefix}{opt_code}. {opt_text}")
                    if is_selected:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green
            
            doc.add_paragraph() # Spacer
        
        # Total Score Summary
        total_p = doc.add_paragraph()
        total_run = total_p.add_run(f"TOTAL ASSESSMENT SCORE: {assessment_data.get('calculated_score', 0)} / 100")
        total_run.bold = True
        total_run.font.size = Pt(11)
        doc.add_paragraph()

        # 5a. Scoring Reference Chart
        doc.add_heading('SCORING REFERENCE', level=1)
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.style = 'Table Grid'
        hdr_cells = ref_table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Scoring Rule (Option: Points)'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].paragraphs[0].runs[0].bold = True

        from app.services.risk_profile_service import SCORING_RULES
        for q_id, rules in SCORING_RULES.items():
            if q_id == 'q2': continue
            row_cells = ref_table.add_row().cells
            row_cells[0].text = q_id.upper()
            row_cells[1].text = ", ".join([f"{opt.upper()}: {score}" for opt, score in rules.items()])
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Q2 Scoring Factors: (A: Very, B: Somewhat, C: Not Important)").bold = True
        p.style.font.size = Pt(9)
        
        q2_table = doc.add_table(rows=1, cols=4)
        q2_table.style = 'Table Grid'
        hdr_cells = q2_table.rows[0].cells
        hdr_cells[0].text = 'Factor'
        hdr_cells[1].text = 'A'
        hdr_cells[2].text = 'B'
        hdr_cells[3].text = 'C'
        for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

        for f_id, f_name in QUESTIONNAIRE_DATA['q2']['factors'].items():
            f_rules = SCORING_RULES['q2'].get(f_id, {})
            row_cells = q2_table.add_row().cells
            row_cells[0].text = f_name
            row_cells[1].text = str(f_rules.get('A', 0))
            row_cells[2].text = str(f_rules.get('B', 0))
            row_cells[3].text = str(f_rules.get('C', 0))
            for cell in row_cells: cell.paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph()

        # 6. Recommendation and Notes
        doc.add_heading('ADVISOR RECOMMENDATION', level=1)
        p = doc.add_paragraph()
        p.add_run("Classification: ").bold = True
        p.add_run(str(assessment_data.get('assigned_risk_tier', 'N/A')))
        
        doc.add_paragraph(assessment_data.get('tier_recommendation') or "No recommendation provided.")
        
        # Add free space for manual recommendation
        doc.add_paragraph().add_run("Additional Advisor Guidance:").bold = True
        for _ in range(8): doc.add_paragraph() # Significant blank space
        
        # DISCUSSION NOTES & DISCLAIMER
        doc.add_page_break()

        if assessment_data.get('discussion_notes'):
            doc.add_heading('DISCUSSION NOTES', level=1)
            p = doc.add_paragraph()
            p.add_run(assessment_data['discussion_notes'])

        if assessment_data.get('disclaimer_text'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("DISCLAIMER: ")
            run.bold = True
            p.add_run(assessment_data['disclaimer_text'])

        # 7. Signature Section
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.width = Inches(6)
        
        cells = sig_table.rows[0].cells
        cells[0].text = "__________________________\nClient Signature"
        cells[1].text = "__________________________\nIA Advisor Signature"
        
        cells = sig_table.rows[1].cells
        cells[0].text = f"Date: ________"
        cells[1].text = f"Date: ________"

        # 8. Save to Buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    @staticmethod
    def generate_risk_profile_pdf_bridge(
        assessment_data: dict,
        client_data: dict,
        ia_data: dict,
        ia_logo_override: str = None,
        questionnaire_data: dict = None
    ) -> BytesIO:
        """PDF generator optimized for Bridge data (dictionaries)."""
        # 1. Prepare data objects (Mock objects for attribute access compatibility)
        class MockObj:
            def __init__(self, d):
                for k, v in d.items():
                    setattr(self, k, v)
            def get(self, k, default=None):
                return getattr(self, k, default)

        assessment = MockObj(assessment_data)
        client = MockObj(client_data)
        ia = MockObj(ia_data)
        
        # Handle Date parsing
        if isinstance(assessment.submitted_at, str):
            try:
                from dateutil import parser
                assessment.submitted_at = parser.parse(assessment.submitted_at)
            except:
                assessment.submitted_at = datetime.now()

        # Handle Questionnaire
        if questionnaire_data:
            questionnaire = MockObj(questionnaire_data)
        else:
            # Fallback or error
            raise ValueError("Questionnaire data required for Bridge PDF generation")

        # 2. Setup Document
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=1.2*inch,
            bottomMargin=0.7*inch, 
            leftMargin=0.5*inch, 
            rightMargin=0.5*inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Define Styles
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#1e3c72'),
            spaceAfter=20, alignment=1, fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2a5298'),
            spaceAfter=12, spaceBefore=15, fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'NormalCustom', parent=styles['Normal'], fontSize=9, spaceAfter=4, fontName='Helvetica'
        )
        bold_style = ParagraphStyle(
            'BoldCustom', parent=styles['Normal'], fontSize=9, spaceAfter=4, fontName='Helvetica-Bold'
        )

        # Cover Page
        resolved_logo = resolve_logo_path(ia_logo_override or (ia.ia_logo_path if ia else None))
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=2.5*inch, height=1.25*inch, kind='proportional')
                story.append(Spacer(1, 100))
                story.append(logo)
                story.append(Spacer(1, 50))
            except: pass
            
        story.append(Paragraph("RISK PROFILE ASSESSMENT REPORT", title_style))
        story.append(Spacer(1, 30))
        story.append(Paragraph(f"<b>CLIENT NAME:</b> {client.client_name}", normal_style))
        story.append(Paragraph(f"<b>CLIENT CODE:</b> {client.client_code}", normal_style))
        if ia:
            story.append(Paragraph(f"<b>ENTITY:</b> {ia.name_of_entity or ia.name_of_ia}", normal_style))
        story.append(Paragraph(f"<b>DATE:</b> {assessment.submitted_at.strftime('%B %d, %Y')}", normal_style))
        
        story.append(PageBreak())

        # Summary Table
        story.append(Paragraph("ASSESSMENT SUMMARY", heading_style))
        summary_data = [
            ["Total Score:", f"{assessment.total_score} / {getattr(questionnaire, 'max_possible_score', 100)}"],
            ["Risk Category:", assessment.category_name or "N/A"]
        ]
        t = Table(summary_data, colWidths=[2*inch, 3*inch])
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
            ('BACKGROUND', (0,0), (0,-1), colors.whitesmoke),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # Responses
        story.append(Paragraph("QUESTIONNAIRE RESPONSES", heading_style))
        responses = getattr(assessment, 'responses', {}) or getattr(assessment, 'answers', {}) or {}
        
        q_list = getattr(questionnaire, 'questions', [])
        for i, q in enumerate(q_list):
            q_id = q.get('id', f'q_{i}')
            # Key might be q_177... (as seen in logs)
            ans = responses.get(q_id, {})
            
            story.append(Paragraph(f"<b>{i+1}. {q.get('text', '')}</b> (Score: {ans.get('score', 0)})", normal_style))
            
            selected_opt_id = ans.get('option_id')
            for opt in q.get('options', []):
                is_sel = str(opt.get('id')) == str(selected_opt_id)
                prefix = "<b>[X]</b> " if is_sel else "[ ] "
                color = colors.darkgreen if is_sel else colors.black
                
                opt_style = ParagraphStyle('Opt', parent=normal_style, textColor=color, leftIndent=20)
                story.append(Paragraph(f"{prefix}{opt.get('text', '')}", opt_style))
            
            story.append(Spacer(1, 10))

        # Scoring Reference
        story.append(PageBreak())
        story.append(Paragraph("SCORING REFERENCE", heading_style))
        cat_data = [["Category", "Score Range", "Description"]]
        for cat in getattr(questionnaire, 'categories', []):
            cat_data.append([
                cat.get('name', ''),
                f"{cat.get('min_score')} - {cat.get('max_score')}",
                cat.get('description', '')
            ])
        
        ct = Table(cat_data, colWidths=[1.5*inch, 1*inch, 3.5*inch])
        ct.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
        ]))
        story.append(ct)

        # Discussion
        if assessment.discussion_notes:
            story.append(Spacer(1, 20))
            story.append(Paragraph("DISCUSSION & NOTES", heading_style))
            story.append(Paragraph(assessment.discussion_notes, normal_style))

        # Signatures
        story.append(Spacer(1, 50))
        sig_data = [
            ["__________________________", "__________________________"],
            ["Client Signature", "Advisor Signature"],
            [f"Date: {assessment.submitted_at.strftime('%Y-%m-%d')}", f"Date: {assessment.submitted_at.strftime('%Y-%m-%d')}"]
        ]
        st = Table(sig_data, colWidths=[3*inch, 3*inch])
        st.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(st)

        doc.build(story, onFirstPage=ReportService._draw_footer, onLaterPages=ReportService._draw_footer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_risk_profile_docx_bridge(
        assessment_data: dict,
        client_data: dict,
        ia_data: dict,
        ia_logo_override: str = None,
        questionnaire_data: dict = None
    ) -> BytesIO:
        """Word generator optimized for Bridge data (dictionaries)."""
        # 1. Prepare data objects
        class MockObj:
            def __init__(self, d):
                for k, v in d.items():
                    setattr(self, k, v)
        
        assessment = MockObj(assessment_data)
        client = MockObj(client_data)
        ia = MockObj(ia_data)
        
        # Handle Date parsing
        if isinstance(assessment.submitted_at, str):
            try:
                from dateutil import parser
                assessment.submitted_at = parser.parse(assessment.submitted_at)
            except:
                assessment.submitted_at = datetime.now()

        # Handle Questionnaire
        if questionnaire_data:
            questionnaire = MockObj(questionnaire_data)
        else:
            raise ValueError("Questionnaire data required for Bridge Word generation")

        doc = Document()
        
        # Header/Footer Setup
        section = doc.sections[0]
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "STRICTLY CONFIDENTIAL"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.style.font.bold = True
        header_p.style.font.size = Pt(8)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Page ")
        ReportService._add_page_number(footer_p.add_run())

        # Cover Page
        resolved_logo = resolve_logo_path(ia_logo_override or (ia.ia_logo_path if ia else None))
        if resolved_logo:
            try:
                doc.add_picture(resolved_logo, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass

        for _ in range(5): doc.add_paragraph()
        title = doc.add_heading('RISK PROFILE ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"\n\n\nCLIENT NAME: {client.client_name}\n").bold = True
        subtitle.add_run(f"CLIENT CODE: {client.client_code}\n")
        if ia:
            entity_name = getattr(ia, 'name_of_entity', None) or getattr(ia, 'name_of_ia', 'N/A')
            subtitle.add_run(f"ENTITY: {entity_name}\n")
        subtitle.add_run(f"DATE: {assessment.submitted_at.strftime('%B %d, %Y')}")
        
        doc.add_page_break()

        # Summary
        doc.add_heading('ASSESSMENT SUMMARY', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        summary_rows = [
            ("Client Name:", client.client_name),
            ("Date of Assessment:", assessment.submitted_at.strftime("%Y-%m-%d %H:%M")),
            ("Total Score:", f"{assessment.total_score} / {getattr(questionnaire, 'max_possible_score', 100)}"),
            ("Risk Category:", assessment.category_name or "N/A")
        ]
        for i, (label, value) in enumerate(summary_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_paragraph().add_run().add_break()

        # Responses
        doc.add_heading('QUESTIONNAIRE RESPONSES', level=1)
        responses = getattr(assessment, 'responses', {}) or getattr(assessment, 'answers', {}) or {}
        q_list = getattr(questionnaire, 'questions', [])
        
        for i, q in enumerate(q_list):
            q_id = q.get('id', f'q_{i}')
            ans = responses.get(q_id, {})
            score_text = f" (Score: {ans.get('score', 0)})"
            
            p = doc.add_paragraph()
            p.add_run(f"{i+1}. {q.get('text', '')}{score_text}").bold = True
            
            selected_opt_id = ans.get('option_id')
            for opt in q.get('options', []):
                is_sel = str(opt.get('id')) == str(selected_opt_id)
                prefix = "[✓] " if is_sel else "[  ] "
                
                option_p = doc.add_paragraph()
                option_p.paragraph_format.left_indent = Pt(20)
                r = option_p.add_run(f"{prefix}{opt.get('text', '')}")
                if is_sel:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0, 100, 0)
            
            doc.add_paragraph()

        # Discussion
        if assessment.discussion_notes:
            doc.add_heading('DISCUSSION & NOTES', level=1)
            doc.add_paragraph(assessment.discussion_notes)

        # Signatures
        doc.add_paragraph().add_run().add_break()
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "__________________________\nClient Signature"
        sig_table.rows[0].cells[1].text = "__________________________\nAdvisor Signature"
        sig_date = assessment.submitted_at.strftime("%Y-%m-%d")
        sig_table.rows[1].cells[0].text = f"Date: {sig_date}"
        sig_table.rows[1].cells[1].text = f"Date: {sig_date}"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
