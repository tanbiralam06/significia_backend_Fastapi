"""
Financial Report Generator — PDF and Word Document Generation.
Adapted from finplan.py with SQLAlchemy support and modern styling.
"""
import os
import io
import re
import uuid
import json
from datetime import datetime
from typing import Optional, List, Dict, Any

# PDF generation imports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
    
    class NumberedCanvas(canvas.Canvas):
        """Custom PDF Canvas class with 'Page x of y' in footer and Entity Name in header."""
        def __init__(self, *args, **kwargs):
            self.entity_name = kwargs.pop('entity_name', "")
            self.advisor_name = kwargs.pop('advisor_name', "")
            self.ia_reg_no = kwargs.pop('ia_reg_no', "")
            canvas.Canvas.__init__(self, *args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_number(num_pages)
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

        def draw_page_number(self, page_count):
            # Footer: Page numbers
            self.setFont("Helvetica", 9)
            current_page = self._pageNumber
            page_text = f"Page {current_page} of {page_count}"
            self.drawRightString(570, 20, page_text)
            
            # Footer: Prepared by / Entity Name / Reg No (left-aligned)
            if any([self.advisor_name, self.entity_name, self.ia_reg_no]):
                self.setFont("Helvetica-Oblique", 7)
                self.setFillColor(colors.grey)
                footer_parts = []
                if self.advisor_name: footer_parts.append(f"Prepared by: {self.advisor_name}")
                if self.entity_name: footer_parts.append(f"Entity: {self.entity_name}")
                if self.ia_reg_no: footer_parts.append(f"Reg No: {self.ia_reg_no}")
                footer_text = " , ".join(footer_parts)
                self.drawString(30, 20, footer_text)
            
            # Header: Entity Name (top-left)
            if self.entity_name:
                self.setFont("Helvetica-Bold", 8)
                self.setFillColor(colors.HexColor('#1e293b'))
                self.drawString(30, 820, self.entity_name.upper())

            # Header: STRICTLY CONFIDENTIAL (top-right, every page)
            self.setFont("Helvetica-Oblique", 7)
            self.setFillColor(colors.grey)
            self.drawRightString(570, 820, "STRICTLY CONFIDENTIAL")
            self.setFillColor(colors.black)

except ImportError:
    PDF_AVAILABLE = False
    
    class NumberedCanvas:
        """Dummy class when reportlab is missing."""
        def __init__(self, *args, **kwargs):
            pass

# Word document generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


def format_currency(val: float) -> str:
    """Format float to currency string with Rs prefix."""
    if val is None:
        return "Rs 0"
    return f"Rs {val:,.0f}"

def format_number(val: float) -> str:
    """Format float to number string with commas."""
    if val is None:
        return "0"
    return f"{val:,.0f}"

def resolve_logo_path(logo_path: Optional[str]) -> Optional[str]:
    """Try multiple strategies to find the logo file on disk."""
    if not logo_path:
        return None
        
    # Strategy 1: Absolute path (most likely from resolve_logo_to_local_path utility)
    if os.path.isabs(logo_path) and os.path.exists(logo_path):
        return logo_path
        
    # Strategy 2: Relative to CWD
    if os.path.exists(logo_path):
        return os.path.abspath(logo_path)

    # Strategy 3: Relative to backend root
    file_dir = os.path.dirname(os.path.abspath(__file__))
    backend_root = os.path.abspath(os.path.join(file_dir, '..', '..'))
    joined_path = os.path.join(backend_root, logo_path)
    if os.path.exists(joined_path):
        return joined_path

    # Strategy 4: Try prepending 'uploads/' if it's a relative path starting with 'ia_documents'
    if not logo_path.startswith('uploads/') and 'ia_documents' in logo_path:
        uploads_path = os.path.join(backend_root, 'uploads', logo_path)
        if os.path.exists(uploads_path):
            return uploads_path

    return None


class FinancialReportGenerator:
    """
    Generate PDF and Word reports with full parity to legacy finplan.py.
    """

    @staticmethod
    def _parse_professional_text(text: str) -> List[Dict[str, str]]:
        """
        Parses unformatted AI text into structured segments.
        Identifies ALL CAPS headers and separates them from paragraphs.
        """
        if not text:
            return []
            
        # Clean up common AI formatting artifacts
        text = text.replace("**", "").replace("__", "")
        
        # Regex to find potential ALL CAPS headers
        # Matches sequences of uppercase letters and spaces (at least 5 chars) 
        # that are followed by a newline, colon, or just standalone at the start/end of a block.
        pattern = r'([A-Z\s]{8,}(?::|\n|$))'
        parts = re.split(pattern, text)
        
        sections = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            # Heuristic for headers: All caps (with some common allowed chars) and length > 5
            is_header = re.match(r'^[A-Z\s\(\)\d\:\-\,]+$', part) and len(part) > 5
            
            # Special case for "Prepared for" or "Analysis Date" - those are meta-data, not section headers
            if is_header and any(meta in part.lower() for meta in ["prepared for", "analysis date"]):
                is_header = False # Treat as normal text or metadata
            
            if is_header:
                sections.append({"type": "header", "content": part.rstrip(':')})
            else:
                # If it's normal text, treat as paragraph
                sections.append({"type": "paragraph", "content": part})
                
        return sections

    @staticmethod
    def generate_pdf(
        result: Any,
        profile: Any,
        client_name: str,
        ia_logo_path: Optional[str] = None,
        ia_name: Optional[str] = None
    ) -> io.BytesIO:
        """Generate a professionally formatted PDF Financial Analysis report."""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is not installed.")

        buffer = io.BytesIO()
        
        # Factory for canvas with ia_name and advisor_name
        def canvas_factory(*args, **kwargs):
            advisor_name = profile.client.advisor_name if hasattr(profile, 'client') and profile.client else None
            ia_reg_no = profile.client.advisor_registration_number if hasattr(profile, 'client') and profile.client else None
            return NumberedCanvas(*args, entity_name=ia_name, advisor_name=advisor_name, ia_reg_no=ia_reg_no, **kwargs)

        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=50, bottomMargin=40)
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('ReportTitle', parent=styles['Heading1'], fontSize=24, alignment=1, spaceAfter=20)
        section_style = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=14, spaceBefore=20, spaceAfter=10, textColor=colors.HexColor('#1e293b'))
        subsection_style = ParagraphStyle('SubSectionHeader', parent=styles['Heading3'], fontSize=11, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold')
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.leading = 14

        # Cover Page
        elements.append(Spacer(1, 100))
        resolved_logo = resolve_logo_path(ia_logo_path)
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=1.5*inch, height=1.5*inch)
                elements.append(logo)
            except Exception as e:
                print(f"Error rendering logo: {e}")
        
        elements.append(Spacer(1, 20))
        elements.append(Paragraph('FINANCIAL ANALYSIS REPORT', title_style))
        elements.append(Spacer(1, 10))
        prepared_by = ia_name or profile.client.advisor_name or 'INVESTMENT ADVISOR'
        cover_details = [
            [Paragraph(f"<b>CLIENT NAME:</b> {client_name.upper()}", normal_style)],
            [Paragraph(f"<b>PREPARED BY:</b> {prepared_by.upper()}", normal_style)],
            [Paragraph(f"<b>REPORT DATE:</b> {datetime.now().strftime('%d %B, %Y').upper()}", normal_style)]
        ]
        if profile.client.client_code:
            cover_details.insert(1, [Paragraph(f"<b>CLIENT CODE:</b> {profile.client.client_code}", normal_style)])
            
        t_cover = Table(cover_details, colWidths=[4.5*inch])
        t_cover.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ]))
        elements.append(t_cover)
        
        elements.append(PageBreak())
        # --- END PREMIUM COVER PAGE ---

        # 1. CLIENT INFORMATION
        elements.append(Paragraph('1. Client Information', section_style))

        # 1.1 Basic Info
        elements.append(Paragraph('1.1 Client Basic Information', subsection_style))
        basic_data = [
            ['Particulars', 'Details'],
            ['Occupation', profile.occupation],
            ['Date of Birth', str(profile.dob)],
            ['Annual Income', format_currency(profile.annual_income)],
            ['Email ID', profile.email or 'Not provided'],
            ['PAN Number', profile.pan or 'Not provided']
        ]
        t = Table(basic_data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('BACKGROUND', (0,1), (-1,-1), colors.white)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.2 Spouse Info
        elements.append(Paragraph('1.2 Spouse Information', subsection_style))
        spouse_data = [
            ['Particulars', 'Details'],
            ['Spouse Name', profile.spouse_name or 'Not provided'],
            ['Date of Birth', str(profile.spouse_dob) if profile.spouse_dob else 'Not provided'],
            ['Occupation', profile.spouse_occupation or 'Not provided']
        ]
        t = Table(spouse_data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('BACKGROUND', (0,1), (-1,-1), colors.white)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.3 Children Info
        if profile.children:
            elements.append(Paragraph('1.3 Children Information', subsection_style))
            child_data = [['Child Name', 'Date of Birth', 'Occupation/Status']]
            for child in profile.children:
                child_data.append([child.get('name', 'N/A'), child.get('dob', 'N/A'), child.get('occupation', 'N/A')])
            t = Table(child_data, colWidths=[150, 150, 200])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('GRID', (0,0), (-1,-1), 1, colors.black)
            ]))
            elements.append(t)
            elements.append(Spacer(1, 12))

        # 1.4 Contact & Identification
        elements.append(Paragraph('1.4 Contact & Identification', subsection_style))
        contact_data = [
            ['Particulars', 'Details'],
            ['Address', profile.client.address or 'Not provided'],
            ['Contact Number', profile.contact]
        ]
        t = Table(contact_data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('BACKGROUND', (0,1), (-1,-1), colors.white)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.5 Expenses Breakdown
        elements.append(Paragraph('1.5 Annual Expenses', subsection_style))
        exp_cats = {
            'hh': 'Household & Groceries', 'med': 'Medical & Healthcare', 'travel': 'Travel & Transportation',
            'elec': 'Electricity & Utilities', 'tele': 'Telephone & Internet', 'maid': 'Maid & Domestic Help',
            'edu': 'Education & Children', 'ent': 'Entertainment & Leisure', 'emi': 'EMI Paid',
            'savings': 'Savings/Investment CONTRIBUTION', 'misc': 'Miscellaneous Expenses'
        }
        exp_data = [['Expense Category', 'Amount (Rs)']]
        total_exp = 0
        for k, v in exp_cats.items():
            amt = profile.expenses.get(k, 0)
            if amt > 0:
                exp_data.append([v, format_number(amt)])
                total_exp += amt
        exp_data.append(['TOTAL ANNUAL EXPENSES', format_number(total_exp)])
        t = Table(exp_data, colWidths=[300, 200])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.6 & 1.7 Assets and Liabilities
        elements.append(Paragraph('1.6 Assets & 1.7 Liabilities', subsection_style))
        ast_data = [['Asset Category', 'Amount (Rs)']]
        ast_map = {'land': 'Land & Building', 'inv': 'Investments', 'cash': 'Cash at Bank', 'retirement': 'Retirement Savings'}
        total_ast = 0
        for k, label in ast_map.items():
            v = profile.assets.get(k, 0)
            if v > 0:
                ast_data.append([label, format_number(v)])
                total_ast += v
        ast_data.append(['TOTAL ASSETS', format_number(total_ast)])
        t = Table(ast_data, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 6))

        lib_data = [['Liability Category', 'Amount (Rs)']]
        lib_map = {'personal': 'Personal Loan', 'cc': 'Credit Card', 'hb': 'Home/Building Loan'}
        total_lib = 0
        # Standard liabilities
        for k, label in lib_map.items():
            v = profile.liabilities.get(k, 0)
            if v > 0:
                lib_data.append([label, format_number(v)])
                total_lib += v
                
        # Custom "Other" liabilities
        others = profile.liabilities.get('others', [])
        if isinstance(others, list):
            for other in others:
                amt = other.get('amount', 0) if isinstance(other, dict) else getattr(other, 'amount', 0)
                if amt > 0:
                    label = other.get('label', 'Other') if isinstance(other, dict) else getattr(other, 'label', 'Other')
                    lib_data.append([label, format_number(amt)])
                    total_lib += amt

        lib_data.append(['TOTAL LIABILITIES', format_number(total_lib)])
        t = Table(lib_data, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.8 Net Worth Calculation
        elements.append(Paragraph('1.8 Net Worth Calculation', subsection_style))
        net_worth = total_ast - total_lib
        nw_data = [
            ['Total Assets', format_number(total_ast)],
            ['Total Liabilities', format_number(total_lib)],
            ['ESTIMATED NET WORTH', format_number(net_worth)]
        ]
        t = Table(nw_data, colWidths=[300, 200])
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 1.9 Current Insurance Cover
        elements.append(Paragraph('1.9 Current Insurance Cover', subsection_style))
        ins_data = [
            ['Insurance Type', 'Cover Amount (Rs)', 'Premium (Rs)'],
            ['Life Insurance', format_number(profile.insurance.get('life_cover', 0)), format_number(profile.insurance.get('life_premium', 0))],
            ['Medical Cover', format_number(profile.insurance.get('medical_cover', 0)), format_number(profile.insurance.get('medical_premium', 0))],
            ['Vehicle insurance', format_number(profile.insurance.get('vehicle_cover', 0)), format_number(profile.insurance.get('vehicle_premium', 0))],
        ]
        t = Table(ins_data, colWidths=[200, 150, 150])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 20))

        # 2. Financial Assumptions
        elements.append(Paragraph('2. Financial Assumptions', section_style))
        ass_data = [['Parameter', 'Value']]
        
        # Label mapping for assumptions
        label_map = {
            'sol_hlv': 'Standard of Living for HLV',
            'sol_ret': 'Standard of Living for Retirement',
            'le_client': 'Life Expectancy - Client',
            'le_spouse': 'Life Expectancy - Spouse',
            'inc_inc_rate': 'Income Increment Rate',
            'inflation': 'Inflation Rate',
            'medical_inflation': 'Medical Inflation Rate',
            'pre_ret_rate': 'Pre-Retirement Return Rate',
            'post_ret_rate': 'Post-Retirement Return Rate',
            'retirement_age': 'Retirement Age',
            'education_years': 'Years to Education Goal',
            'marriage_years': 'Years to Marriage Goal',
            'child_education_corpus': 'Child Education Corpus',
            'child_marriage_corpus': 'Child Marriage Corpus'
        }
        
        for k, v in profile.assumptions.items():
            label = label_map.get(k, k.replace('_', ' ').title())
            if 'corpus' in k:
                val_str = format_currency(v)
            elif 'rate' in k or 'inflation' in k or 'sol' in k:
                val_str = f"{v}%"
            else:
                val_str = str(int(v)) if float(v).is_integer() else str(v)
            ass_data.append([label, val_str])
            
        t = Table(ass_data, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Paragraph("", normal_style))

        # 3. Executive Brief
        if not profile.exclude_ai and result.ai_analysis and 'executive_brief' in result.ai_analysis:
            elements.append(Paragraph('Executive Summary & Insights', section_style))
            brief = result.ai_analysis['executive_brief']
            # Remove any HTML tags if they exist
            brief = re.sub(r'<[^>]*>', '', brief)
            
            sections = FinancialReportGenerator._parse_professional_text(brief)
            for sec in sections:
                if sec['type'] == 'header':
                    elements.append(Paragraph(sec['content'], subsection_style))
                else:
                    # Handle internal line breaks within paragraphs
                    para_text = sec['content'].replace('\n', '<br/>')
                    elements.append(Paragraph(para_text, normal_style))
                    elements.append(Spacer(1, 6))
            
            elements.append(Spacer(1, 12))

        # 4. HLV
        elements.append(Paragraph('4. Human Life Value Analysis', section_style))
        hlv = result.hlv_data
        hlv_info = [
            ['Description', 'Value (Rs)'],
            ['HLV (Income Replacement Method)', format_number(hlv.get('hlv_income_method'))],
            ['HLV (Need Based Method)', format_number(hlv.get('hlv_expense_method'))],
            ['Net HLV (Income)', format_number(hlv.get('net_hlv_income'))],
            ['Additional Life Cover Needed (Income)', format_number(hlv.get('additional_life_cover_needed_income'))],
            ['Current Life Cover', format_number(profile.insurance.get('life_cover', 0))],
        ]
        t = Table(hlv_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        #HLV Comments
        if not profile.exclude_ai and result.ai_analysis and 'hlv_comments' in result.ai_analysis:
            elements.append(Paragraph('Insights (HLV):', subsection_style))
            for comment in result.ai_analysis['hlv_comments']:
                clean_comment = comment.replace('AI Insight', 'Insight').replace('<strong>Insight', '<strong>Insight')
                elements.append(Paragraph(f"• {clean_comment}", normal_style))
            elements.append(Spacer(1, 6))

        # 5. Medical
        elements.append(Paragraph('5. Medical Coverage Analysis', section_style))
        med = result.medical_data
        med_info = [
            ['Description', 'Value (Rs)'],
            ['Medical Corpus at Retirement', format_number(med.get('medical_corpus_at_retirement'))],
            ['Medical Corpus at Life Expectancy', format_number(med.get('medical_corpus_at_life_expectancy'))],
            ['Balance Needed at Retirement', format_number(med.get('balance_needed_at_retirement'))],
        ]
        t = Table(med_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # AI Medical Comments
        if not profile.exclude_ai and result.ai_analysis and 'medical_comments' in result.ai_analysis:
            elements.append(Paragraph('Insights (Medical):', subsection_style))
            for comment in result.ai_analysis['medical_comments']:
                clean_comment = comment.replace('AI Insight', 'Insight')
                elements.append(Paragraph(f"• {clean_comment}", normal_style))
            elements.append(Spacer(1, 6))

        # 6. Retirement
        elements.append(Paragraph('6. Retirement Corpus Analysis', section_style))
        ret = result.calculations
        ret_info = [
            ['Description', 'Value (Rs)'],
            ['Retirement Corpus Needed', format_number(ret.get('retirement_corpus_at_retirement'))],
            ['Net Corpus Needed', format_number(ret.get('net_retirement_corpus_needed'))],
            ['Retirement Readiness', f"{ret.get('retirement_readiness', 0)}%"],
        ]
        t = Table(ret_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # AI Retirement Comments
        if not profile.exclude_ai and result.ai_analysis and 'retirement_comments' in result.ai_analysis:
            elements.append(Paragraph('Insights (Retirement):', subsection_style))
            for comment in result.ai_analysis['retirement_comments']:
                clean_comment = comment.replace('AI Insight', 'Insight')
                elements.append(Paragraph(f"• {clean_comment}", normal_style))
            elements.append(Spacer(1, 6))

        # 7. Cash Flow Table
        if result.cash_flow_analysis:
            elements.append(Paragraph('7. Retirement Cash Flow Analysis', section_style))
            cf = result.cash_flow_analysis
            cf_data = [['Year', 'Age', 'Opening Bal (Rs)', 'Growth (Rs)', 'Withdrawal (Rs)', 'Closing Bal (Rs)']]
            for row in result.cash_flow_analysis:
                cf_data.append([
                    str(row['year']), str(row['retirement_age_year']),
                    format_number(row['opening_balance']), format_number(row['investment_growth']),
                    format_number(row['annual_withdrawal']), format_number(row['closing_balance'])
                ])
            
            t = Table(cf_data, colWidths=[40, 40, 100, 100, 100, 100])
            t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('FONTSIZE', (0,0), (-1,-1), 8), ('GRID', (0,0), (-1,-1), 0.5, colors.black)]))
            elements.append(t)

        # 8. Child Goals
        elements.append(Paragraph('8. Child Goals Analysis', section_style))
        edu_info = [
            ['Education Today Value (Rs)', format_number(result.calculations.get('child_education_corpus_today', 0))],
            ['Education Future Needed (Rs)', format_number(result.calculations.get('child_education_future_needed', 0))],
            ['Education Net Corpus (Rs)', format_number(result.calculations.get('education_net_corpus', 0))],
            ['Education Monthly Investment (Rs)', format_number(result.calculations.get('monthly_investment_education', 0))],
        ]
        t = Table(edu_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 6))

        marr_info = [
            ['Marriage Today Value (Rs)', format_number(result.calculations.get('marriage_corpus_today', 0))],
            ['Marriage Future Needed (Rs)', format_number(result.calculations.get('marriage_future_needed', 0))],
            ['Marriage Net Corpus (Rs)', format_number(result.calculations.get('marriage_net_corpus', 0))],
            ['Marriage Monthly Investment (Rs)', format_number(result.calculations.get('monthly_investment_marriage', 0))],
        ]
        t = Table(marr_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 9. Emergency Fund Analysis
        elements.append(Paragraph('9. Emergency Fund Analysis', section_style))
        em_info = [
            ['Emergency Fund Needed (6 months) (Rs)', format_number(result.calculations.get('emergency_fund_needed', 0))],
            ['Shortfall (Rs)', format_number(result.calculations.get('emergency_fund_shortfall', 0))],
            ['Monthly Investment Required (Rs)', format_number(result.calculations.get('monthly_investment_emergency', 0))],
        ]
        t = Table(em_info, colWidths=[300, 200])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 10. Summary Table
        elements.append(Paragraph('10. Monthly Investment Summary', section_style))
        inv_data = [
            ['Investment Goal', 'Monthly Investment', 'Months'],
            ['Retirement', format_currency(ret.get('monthly_investment_retirement')), str(int(ret.get('years_to_retirement', 0)*12))],
            ['Education', format_currency(ret.get('monthly_investment_education')), str(int(profile.assumptions.get('education_years', 0)*12))],
            ['Marriage', format_currency(ret.get('monthly_investment_marriage')), str(int(profile.assumptions.get('marriage_years', 0)*12))],
            ['TOTAL (Income Method)', format_currency(ret.get('total_monthly_investment_income')), ''],
            ['TOTAL (Expense Method)', format_currency(ret.get('total_monthly_investment_expense')), '']
        ]
        t = Table(inv_data, colWidths=[200, 150, 100])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('FONTNAME', (0,-2), (-1,-1), 'Helvetica-Bold'), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)

        # 11. Investment Allocation
        elements.append(Paragraph('11. Existing Investment Allocation(%)', section_style))
        all_data = [
            ['Goal', 'Allocation Percentage'],
            ['Education Goal', f"{profile.education_investment_pct}%"],
            ['Marriage Goal', f"{profile.marriage_investment_pct}%"],
        ]
        t = Table(all_data, colWidths=[250, 250])
        t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.grey), ('GRID', (0,0), (-1,-1), 1, colors.black)]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 12. Financial Health Score
        elements.append(Paragraph('12. Financial Health Score', section_style))
        score_data = [['Component', 'Status/Value', 'Points Awarded']]
        for component, status, points in result.calculations.get('financial_health_score_details', []):
            score_data.append([component, str(status), str(points)])
        score_data.append(['TOTAL SCORE', '', f"{result.financial_health_score}/100"])
        t = Table(score_data, colWidths=[200, 200, 100])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # 13. Conclusion
        if not profile.exclude_ai and result.ai_analysis and 'overall_conclusion' in result.ai_analysis:
            elements.append(Paragraph('13. OVERALL CONCLUSION', section_style))
            conc = result.ai_analysis['overall_conclusion']
            conc = re.sub(r'<[^>]*>', '', conc)
            
            sections = FinancialReportGenerator._parse_professional_text(conc)
            for sec in sections:
                if sec['type'] == 'header':
                    elements.append(Paragraph(sec['content'], subsection_style))
                else:
                    para_text = sec['content'].replace('\n', '<br/>')
                    elements.append(Paragraph(para_text, normal_style))
                    elements.append(Spacer(1, 6))
            
            elements.append(Spacer(1, 20))

        # 14. Disclaimer
        elements.append(Paragraph("", normal_style))
        elements.append(Paragraph('14. DISCLAIMER', section_style))
        disc = profile.disclaimer_text or "This report is generated based on data provided by the client..."
        elements.append(Paragraph(disc, normal_style))
        elements.append(Spacer(1, 40))

        # 15. Discussion Notes
        elements.append(Paragraph('15. DISCUSSION NOTES', section_style))
        if profile.discussion_notes:
            elements.append(Paragraph(profile.discussion_notes, normal_style))
        else:
            # Add a large blank space for manual notes
            elements.append(Spacer(1, 200))
        elements.append(Spacer(1, 40))

        # 16. Signatures (Very last)
        elements.append(Spacer(1, 10))
        elements.append(Paragraph('16. SIGNATURES', section_style))
        elements.append(Spacer(1, 40))
        sig_data = [
            ["__________________________", "__________________________"],
            ["Signature of Client", "Signature of IA"],
            [f"Date: {datetime.now().strftime('%d %B, %Y')}", f"Date: {datetime.now().strftime('%d %B, %Y')}"]
        ]
        sig_table = Table(sig_data, colWidths=[250, 250])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 20),
        ]))
        elements.append(sig_table)

        doc.build(elements, canvasmaker=canvas_factory)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_blank_form(ia_logo_path: Optional[str] = None, ia_name: Optional[str] = None, advisor_name: Optional[str] = None, ia_reg_no: Optional[str] = None) -> io.BytesIO:
        """Generate a professionally styled, grid-based blank Financial Analysis Form."""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is not installed.")

        buffer = io.BytesIO()
        # Compact margins for maximum writing space
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        elements = []
        styles = getSampleStyleSheet()

        # Custom Styles
        title_style = ParagraphStyle('FormTitle', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=8, textColor=colors.HexColor('#1e293b'), fontName='Helvetica-Bold')
        section_style = ParagraphStyle('FormSection', parent=styles['Normal'], fontSize=11, textColor=colors.white, fontName='Helvetica-Bold')
        label_style = ParagraphStyle('FormLabel', parent=styles['Normal'], fontSize=9, fontName='Helvetica-Bold', leading=10)
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=7, textColor=colors.grey, alignment=1)

        def add_section_header(text):
            data = [[Paragraph(text.upper(), section_style)]]
            t = Table(data, colWidths=[535]) # Full width minus 30*2 margins
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1e293b')),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 2))

        # --- HEADER ---
        header_data = []
        resolved_logo = resolve_logo_path(ia_logo_path)
        if resolved_logo:
            try:
                logo = Image(resolved_logo, width=0.8*inch, height=0.8*inch)
                header_data.append([logo, Paragraph("FINANCIAL ANALYSIS DATA ENTRY FORM", title_style), ""])
            except Exception as e:
                print(f"Error rendering logo in Blank Form: {e}")
                header_data.append(["", Paragraph("FINANCIAL ANALYSIS DATA ENTRY FORM", title_style), ""])
        else:
            header_data.append(["", Paragraph("FINANCIAL ANALYSIS DATA ENTRY FORM", title_style), ""])
        
        t_header = Table(header_data, colWidths=[80, 375, 80])
        t_header.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('ALIGN', (1,0), (1,0), 'CENTER')]))
        elements.append(t_header)
        elements.append(Spacer(1, 10))

        # Advisor Info Grid
        info_data = [
            [Paragraph("Advisor Name", label_style), "", Paragraph("Advisor ID", label_style), ""],
            [Paragraph("Client Name", label_style), "", Paragraph("Date", label_style), datetime.now().strftime('%d/%m/%Y')],
            [Paragraph("Client Code", label_style), "", "", ""]
        ]
        t_info = Table(info_data, colWidths=[110, 230, 90, 105], rowHeights=25)
        t_info.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.white),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5)
        ]))
        elements.append(t_info)
        elements.append(Spacer(1, 10))

        # 1. PERSONAL INFORMATION
        add_section_header("1. Personal & Family Profile")
        personal_data = [
            [Paragraph("Full Name (Client)", label_style), "", Paragraph("DOB", label_style), ""],
            [Paragraph("Occupation", label_style), "", Paragraph("PAN", label_style), ""],
            [Paragraph("Email ID", label_style), "", Paragraph("Contact", label_style), ""],
            [Paragraph("Annual Income", label_style), "Rs.", "", ""],
            [Paragraph("Full Name (Spouse)", label_style), "", Paragraph("DOB", label_style), ""],
            [Paragraph("Occupation", label_style), "", "", ""]
        ]
        t_personal = Table(personal_data, colWidths=[120, 255, 50, 110], rowHeights=28)
        t_personal.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('SPAN', (1,0), (1,0)), ('SPAN', (1,1), (1,1)), ('SPAN', (1,2), (1,2)),
            ('SPAN', (1,4), (1,4)), ('SPAN', (1,5), (3,5)),
            ('PADDING', (0,0), (-1,-1), 6)
        ]))
        elements.append(t_personal)
        elements.append(Spacer(1, 10))

        # 2. CHILDREN
        add_section_header("2. Children Details")
        child_data = [
            [Paragraph("No.", label_style), Paragraph("Child Full Name", label_style), Paragraph("Date of Birth", label_style), Paragraph("Occupation / Status", label_style)],
            ["1", "", "", ""],
            ["2", "", "", ""],
            ["3", "", "", ""],
        ]
        t_child = Table(child_data, colWidths=[40, 240, 100, 155], rowHeights=30)
        t_child.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (0,-1), 'CENTER'),
            ('PADDING', (0,0), (-1,-1), 8)
        ]))
        elements.append(t_child)
        elements.append(Spacer(1, 10))

        # 3. ANNUAL EXPENSES
        add_section_header("3. Annual Expenses (Estimated Rs.)")
        exp_list = [
            ['Household / Groceries', 'Education & Children'],
            ['Medical & Healthcare', 'Travel & Transport'],
            ['Electricity & Water', 'Telephone & Internet'],
            ['Maid & Domestic Help', 'EMI (Loans)'],
            ['Leisure & Entertainment', 'Insurance Premiums'],
            ['Rent / Maintenance', 'Others Outgo']
        ]
        exp_rows = []
        for pair in exp_list:
            row = []
            for item in pair:
                row.append(Paragraph(item, label_style))
                row.append("Rs. ")
            exp_rows.append(row)
        
        t_exp = Table(exp_rows, colWidths=[160, 107.5, 160, 107.5], rowHeights=30)
        t_exp.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 8)
        ]))
        elements.append(t_exp)

        # 4. ASSETS & 5. LIABILITIES
        elements.append(PageBreak())
        add_section_header("4. Assets & 5. Liabilities (Rs.)")
        al_data = [
            [Paragraph("<b>ASSET DESCRIPTION</b>", label_style), Paragraph("<b>VALUE (Rs.)</b>", label_style), Paragraph("<b>LIABILITY DESCRIPTION</b>", label_style), Paragraph("<b>OUTSTANDING (Rs.)</b>", label_style)],
            ["Real Estate (Res)", "", "Home / Housing Loan", ""],
            ["Real Estate (Other)", "", "Personal / Top-up Loan", ""],
            ["Mutual Funds / Equity", "", "Credit Card Dues", ""],
            ["Fixed Dep / Bank Bal", "", "Vehicle Loan", ""],
            ["Retirement (PF/PPF)", "", "Others LIAB", ""],
            ["Others ASSETS", "", "", ""],
        ]
        t_al = Table(al_data, colWidths=[175, 92.5, 175, 92.5], rowHeights=32)
        t_al.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 8)
        ]))
        elements.append(t_al)
        elements.append(Spacer(1, 10))

        # 6. INSURANCE
        add_section_header("6. Insurance Information")
        ins_data = [
            [Paragraph("Insurance Type", label_style), Paragraph("Sum Assured / Cover", label_style), Paragraph("Annual Premium Paid", label_style)],
            ["Life Insurance (Term/Endow)", "Rs. ", "Rs. "],
            ["Health Insurance (Mediclaim)", "Rs. ", "Rs. "],
            ["Critical Illness / PA", "Rs. ", "Rs. "],
            ["Motor / Asset Insurance", "Rs. ", "Rs. "],
        ]
        t_ins = Table(ins_data, colWidths=[185, 175, 175], rowHeights=32)
        t_ins.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 8)
        ]))
        elements.append(t_ins)
        
        elements.append(Spacer(1, 5))
        mb_data = [[Paragraph("<b>Medical Bonus Years:</b> ", label_style), Paragraph("<b>Avg. Bonus Percentage (%):</b> ", label_style)]]
        t_mb = Table(mb_data, colWidths=[267.5, 267.5], rowHeights=28)
        t_mb.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('PADDING', (0,0), (-1,-1), 6)]))
        elements.append(t_mb)
        elements.append(Spacer(1, 10))

        # 7. ASSUMPTIONS & GOALS
        add_section_header("7. Financial Assumptions & Major Goals")
        ass_data = [
            [Paragraph("Retirement Age", label_style), "", Paragraph("Inflation Rate (%)", label_style), ""],
            [Paragraph("Life Exp (Client)", label_style), "", Paragraph("Life Exp (Spouse)", label_style), ""],
            [Paragraph("Med. Inflation (%)", label_style), "", Paragraph("Pre-Ret Return (%)", label_style), ""],
            [Paragraph("Income Increment (%)", label_style), "", Paragraph("Post-Ret Return (%)", label_style), ""],
            [Paragraph("SOL for HLV (%)", label_style), "", Paragraph("SOL for Retire (%)", label_style), ""],
        ]
        t_ass = Table(ass_data, colWidths=[160, 107.5, 160, 107.5], rowHeights=28)
        t_ass.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 6)
        ]))
        elements.append(t_ass)
        elements.append(Spacer(1, 10))

        goal_data = [
            ["", Paragraph("<b>Education Goal (Children)</b>", label_style), Paragraph("<b>Marriage Goal (Children)</b>", label_style)],
            [Paragraph("Corpus Needed", label_style), "Rs. ", "Rs. "],
            [Paragraph("Years to Goal", label_style), "", ""],
            [Paragraph("Allocation (%)", label_style), "", ""],
        ]
        t_goals = Table(goal_data, colWidths=[145, 195, 195], rowHeights=[25, 32, 32, 32])
        t_goals.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (1,0), (2,0), colors.HexColor('#f1f5f9')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 8)
        ]))
        elements.append(t_goals)

        # 8. NOTES & SIGNATURES
        elements.append(Spacer(1, 15))
        add_section_header("8. Discussion Notes & Signatures")
        notes_data = [[Paragraph("Detailed Discussion Remarks / Client Requests:", label_style)], ["\n\n\n\n\n\n\n\n\n\n"]]
        t_notes = Table(notes_data, colWidths=[535], rowHeights=[25, 300])
        t_notes.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('VALIGN', (0,0), (-1,-1), 'TOP'), ('PADDING', (0,0), (-1,-1), 8)]))
        elements.append(t_notes)
        elements.append(Spacer(1, 30))

        sig_data = [
            ["_________________________", "_________________________"],
            [Paragraph("<b>Client Signature</b>", label_style), Paragraph("<b>Advisor Signature</b>", label_style)],
            ["Name: __________________", "Name: __________________"],
            ["Date: ____ / ____ / 20____", "Date: ____ / ____ / 20____"]
        ]
        t_sig = Table(sig_data, colWidths=[2.8*inch, 2.8*inch], rowHeights=[30, 20, 30, 30])
        t_sig.setStyle(TableStyle([('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
        elements.append(t_sig)

        def canvas_factory(*args, **kwargs):
            return NumberedCanvas(*args, entity_name=ia_name, advisor_name=advisor_name, ia_reg_no=ia_reg_no, **kwargs)

        doc.build(elements, canvasmaker=canvas_factory)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_docx(
        result: Any,
        profile: Any,
        client_name: str,
        ia_logo_path: Optional[str] = None,
        ia_name: Optional[str] = None
    ) -> io.BytesIO:
        """Generate a professionally formatted Word Financial Analysis report."""
        if not WORD_AVAILABLE:
            raise ImportError("python-docx is not installed.")

        doc = Document()
        
        # Add Header with Entity Name and "Strictly Confidential"
        section = doc.sections[0]
        header = section.header
        htable = header.add_table(1, 2, width=Inches(6))
        
        # Left side: IA Name
        if ia_name:
            htable.cell(0, 0).text = ia_name.upper()
            htable.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(8)
            htable.cell(0, 0).paragraphs[0].runs[0].bold = True
            
        # Right side: Strictly Confidential
        htable.cell(0, 1).text = "STRICTLY CONFIDENTIAL"
        htable.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        htable.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(8)
        htable.cell(0, 1).paragraphs[0].runs[0].italic = True

        # Cover Page
        for _ in range(5): doc.add_paragraph()
        
        resolved_logo = resolve_logo_path(ia_logo_path)
        if resolved_logo:
            try:
                doc.add_picture(resolved_logo, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error rendering logo in Word report: {e}")
        
        for _ in range(2): doc.add_paragraph()
        
        title = doc.add_heading("COMPREHENSIVE FINANCIAL ANALYSIS REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(4): doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CLIENT NAME: {client_name.upper()}\n")
        run.bold = True
        if profile.client.client_code:
            p.add_run(f"CLIENT CODE: {profile.client.client_code}\n")
        
        prepared_by = ia_name or profile.client.advisor_name or 'INVESTMENT ADVISOR'
        p.add_run(f"PREPARED BY: {prepared_by}\n")
        p.add_run(f"REPORT DATE: {datetime.now().strftime('%d %B, %Y').upper()}")
        
        doc.add_page_break()
        # --- END PREMIUM COVER PAGE ---

        def add_table(title, data):
            doc.add_heading(title, level=1)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            table.style = 'Table Grid'
            for i, row in enumerate(data):
                for j, val in enumerate(row):
                    table.cell(i, j).text = str(val)

        # 1. Client Profile
        add_table("1.1 Basic Information", [
            ["Particulars", "Details"],
            ["Name", client_name],
            ["Date of Birth", str(profile.dob)],
            ["Occupation", profile.occupation],
            ["Annual Income", format_currency(profile.annual_income)]
        ])

        if profile.spouse_name:
            add_table("1.2 Spouse Information", [
                ["Particulars", "Details"],
                ["Name", profile.spouse_name],
                ["Occupation", profile.spouse_occupation or "N/A"]
            ])

        # 1.5 Expenses
        exp_cats = {
            'hh': 'Household & Groceries', 'med': 'Medical & Healthcare', 'travel': 'Travel & Transportation',
            'elec': 'Electricity & Utilities', 'tele': 'Telephone & Internet', 'maid': 'Maid & Domestic Help',
            'edu': 'Education & Children', 'ent': 'Entertainment & Leisure', 'emi': 'EMI Paid',
            'savings': 'Savings/Investment CONTRIBUTION', 'misc': 'Miscellaneous Expenses'
        }
        exp_data = [["Expense Category", "Amount"]]
        total_exp = 0
        for k, label in exp_cats.items():
            amt = profile.expenses.get(k, 0)
            if amt > 0:
                exp_data.append([label, format_currency(amt)])
                total_exp += amt
        exp_data.append(["TOTAL ANNUAL EXPENSES", format_currency(total_exp)])
        add_table("1.5 Annual Expenses", exp_data)

        # 1.6 Assets
        ast_map = {'land': 'Land & Building', 'inv': 'Investments', 'cash': 'Cash at Bank', 'retirement': 'Retirement Savings'}
        ast_data = [["Asset Category", "Amount"]]
        total_ast = 0
        for k, label in ast_map.items():
            v = profile.assets.get(k, 0)
            if v > 0:
                ast_data.append([label, format_currency(v)])
                total_ast += v
        ast_data.append(["TOTAL ASSETS", format_currency(total_ast)])
        add_table("1.6 Assets", ast_data)

        # 1.7 Liabilities
        lib_map = {'personal': 'Personal Loan', 'cc': 'Credit Card', 'hb': 'Home/Building Loan'}
        lib_data = [["Liability Category", "Amount"]]
        total_lib = 0
        for k, label in lib_map.items():
            v = profile.liabilities.get(k, 0)
            if v > 0:
                lib_data.append([label, format_currency(v)])
                total_lib += v
        others = profile.liabilities.get('others', [])
        if isinstance(others, list):
            for other in others:
                amt = other.get('amount', 0) if isinstance(other, dict) else getattr(other, 'amount', 0)
                if amt > 0:
                    label = other.get('label', 'Other') if isinstance(other, dict) else getattr(other, 'label', 'Other')
                    lib_data.append([label, format_currency(amt)])
                    total_lib += amt
        add_table("1.7 Liabilities", lib_data) # Added this line, it was missing in original code.

        # 2. Financial Assumptions
        label_map = {
            'sol_hlv': 'Standard of Living for HLV',
            'sol_ret': 'Standard of Living for Retirement',
            'le_client': 'Life Expectancy - Client',
            'le_spouse': 'Life Expectancy - Spouse',
            'inc_inc_rate': 'Income Increment Rate',
            'inflation': 'Inflation Rate',
            'medical_inflation': 'Medical Inflation Rate',
            'pre_ret_rate': 'Pre-Retirement Return Rate',
            'post_ret_rate': 'Post-Retirement Return Rate',
            'retirement_age': 'Retirement Age',
            'education_years': 'Years to Education Goal',
            'marriage_years': 'Years to Marriage Goal',
            'child_education_corpus': 'Child Education Corpus',
            'child_marriage_corpus': 'Child Marriage Corpus'
        }
        ass_data = [["Parameter", "Value"]]
        for k, v in profile.assumptions.items():
            label = label_map.get(k, k.replace('_', ' ').title())
            if 'corpus' in k:
                val_str = format_currency(v)
            elif 'rate' in k or 'inflation' in k or 'sol' in k:
                val_str = f"{v}%"
            else:
                val_str = str(int(v)) if float(v).is_integer() else str(v)
            ass_data.append([label, val_str])
        add_table("2. Financial Assumptions", ass_data)

        # 3. Net Worth Summary
        add_table("3. Net Worth Summary", [
            ["Particulars", "Value"],
            ["Total Assets", format_currency(total_ast)],
            ["Total Liabilities", format_currency(total_lib)],
            ["NET WORTH", format_currency(total_ast - total_lib)]
        ])

        # 4. Executive Summary & Insights
        if not profile.exclude_ai and result.ai_analysis and 'executive_brief' in result.ai_analysis:
            doc.add_heading("4. Executive Summary & Insights", level=1)
            brief = result.ai_analysis['executive_brief']
            brief = re.sub(r'<[^>]*>', '', brief)
            sections = FinancialReportGenerator._parse_professional_text(brief)
            for sec in sections:
                if sec['type'] == 'header':
                    doc.add_heading(sec['content'], level=2)
                else:
                    doc.add_paragraph(sec['content'])
            doc.add_page_break()

        # 5. Insurance
        ins = profile.insurance
        add_table("5. Insurance Coverage", [
            ["Type", "Coverage", "Premium"],
            ["Life", format_currency(ins.get('life_cover')), format_currency(ins.get('life_premium'))],
            ["Medical", format_currency(ins.get('medical_cover')), format_currency(ins.get('medical_premium'))],
            ["Vehicle", format_currency(ins.get('vehicle_cover')), format_currency(ins.get('vehicle_premium'))]
        ])

        # 6. HLV
        hlv = result.hlv_data
        hlv_info = [
            ["Description", "Value (Rs)"],
            ["HLV (Income Replacement Method)", format_number(hlv.get('hlv_income_method'))],
            ["HLV (Need Based Method)", format_number(hlv.get('hlv_expense_method'))],
            ["Net HLV (Income)", format_number(hlv.get('net_hlv_income'))],
            ["Additional Life Cover Needed (Income)", format_number(hlv.get('additional_life_cover_needed_income'))],
            ["Current Life Cover", format_number(profile.insurance.get('life_cover', 0))],
        ]
        add_table("6. Human Life Value Analysis", hlv_info)

        # 7/8. Retirement & Medical
        ret = result.calculations
        med = result.medical_data
        rm_info = [
            ["Description", "Value (Rs)"],
            ["Retirement Corpus Needed", format_number(ret.get('retirement_corpus_at_retirement'))],
            ["Net Corpus Needed", format_number(ret.get('net_retirement_corpus_needed'))],
            ["Medical Corpus at Retirement", format_number(med.get('medical_corpus_at_retirement'))],
            ["Retirement Readiness", f"{ret.get('retirement_readiness', 0)}%"]
        ]
        add_table("7/8. Retirement & Medical Analysis", rm_info)

        # 9. Cash Flow
        if result.cash_flow_analysis:
            cf_data = [["Year", "Age", "Opening Bal (Rs)", "Growth (Rs)", "Withdrawal (Rs)", "Closing Bal (Rs)"]]
            for row in result.cash_flow_analysis:
                cf_data.append([str(row['year']), str(row['retirement_age_year']), format_number(row['opening_balance']), 
                                format_number(row['investment_growth']), format_number(row['annual_withdrawal']), format_number(row['closing_balance'])])
            add_table("9. Retirement Cash Flow Analysis", cf_data)

        # 10. Monthly Investment Summary
        sum_data = [
            ['Goal', 'Monthly Investment (Rs)'],
            ['Retirement', format_number(ret.get('monthly_investment_retirement', 0))],
            ['Education', format_number(ret.get('monthly_investment_education', 0))],
            ['Marriage', format_number(ret.get('monthly_investment_marriage', 0))],
            ['TOTAL (Income Method)', format_number(ret.get('total_monthly_investment_income', 0))],
            ['TOTAL (Expense Method)', format_number(ret.get('total_monthly_investment_expense', 0))],
        ]
        add_table("10. Monthly Investment Summary", sum_data)

        # 12. Health Score
        score_details = result.calculations.get('financial_health_score_details', [])
        if score_details:
            score_data = [["Component", "Status", "Points"]]
            for comp, stat, pts in score_details:
                score_data.append([comp, str(stat), str(pts)])
            score_data.append(["TOTAL SCORE", "", f"{result.financial_health_score}/100"])
            add_table("12. Financial Health Score", score_data)

        # 13. Overall Conclusion
        if not profile.exclude_ai and result.ai_analysis and 'overall_conclusion' in result.ai_analysis:
            doc.add_heading("13. OVERALL CONCLUSION", level=1)
            conc = result.ai_analysis['overall_conclusion']
            conc = re.sub(r'<[^>]*>', '', conc)
            
            sections = FinancialReportGenerator._parse_professional_text(conc)
            for sec in sections:
                if sec['type'] == 'header':
                    doc.add_heading(sec['content'], level=2)
                else:
                    doc.add_paragraph(sec['content'])
            doc.add_page_break()

        # 14. Disclaimer
        doc.add_heading("14. DISCLAIMER", level=1)
        disc = profile.disclaimer_text or "This report is generated based on data provided by the client..."
        doc.add_paragraph(disc)
        doc.add_paragraph()
        
        # 15. Discussion Notes
        doc.add_heading("15. DISCUSSION NOTES", level=1)
        if profile.discussion_notes:
            doc.add_paragraph(profile.discussion_notes)
        else:
            for _ in range(10): doc.add_paragraph("__________________________________________________________________")
        
        doc.add_paragraph()
        
        # 16. Signatures
        doc.add_heading("16. Signatures", level=1)
        doc.add_paragraph()
        doc.add_paragraph("__________________________            __________________________")
        doc.add_paragraph("Signature of Client                   Signature of IA")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B, %Y')}            Date: {datetime.now().strftime('%d %B, %Y')}")
        doc.add_paragraph()

        # Add Page Footer
        section = doc.sections[0]
        footer = section.footer
        f_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resolve data
        prepared_by = ia_name or (profile.client.advisor_name if hasattr(profile, 'client') and profile.client else 'INVESTMENT ADVISOR')
        ia_reg_no = profile.client.advisor_registration_number if hasattr(profile, 'client') and profile.client else 'N/A'
        
        footer_parts = [f"Prepared by: {prepared_by}", f"Entity: {ia_name or 'N/A'}", f"Reg No: {ia_reg_no}"]
        footer_text = " | ".join(footer_parts)
        
        f_run = f_p.add_run(footer_text)
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        f_run.italic = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
