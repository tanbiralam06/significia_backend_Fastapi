import io
import os
import base64
from datetime import datetime
from typing import Optional, List

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.ia_master import IAMaster
from app.models.asset_allocation import AssetAllocation

DEFAULT_ASSET_ALLOCATION_DISCLAIMER = """This asset allocation report is prepared based on the client's risk profile and financial goals as assessed by the Investment Advisor. The allocation percentages represent a recommended distribution of investable assets and should be reviewed periodically. Past performance is not indicative of future results. This report does not constitute investment advice and is prepared solely for informational and planning purposes in accordance with SEBI Investment Advisor Regulations."""

class AssetAllocationReportUtils:
    @staticmethod
    def create_pie_chart(labels: List[str], sizes: List[float], title: str, custom_colors: Optional[List[str]] = None) -> bytes:
        """Create a pie chart and return as bytes"""
        if custom_colors is None:
            custom_colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F']
        
        plt.figure(figsize=(4, 3))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=custom_colors[:len(labels)])
        plt.axis('equal')
        plt.title(title, fontsize=10, fontweight='bold', pad=10)
        
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=100, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        
        return buffer.getvalue()

    @staticmethod
    def add_page_number(canvas, doc):
        """Add unified footer and page number to each page"""
        canvas.saveState()
        canvas.setFont("Helvetica-Oblique", 7)
        canvas.setFillColor(colors.grey)
        
        # Resolve data from doc object
        advisor_name = getattr(doc, 'advisor_name', '')
        entity_name = getattr(doc, 'entity_name', '')
        ia_reg_no = getattr(doc, 'ia_reg_no', '')
        
        footer_parts = []
        if advisor_name: footer_parts.append(f"Prepared by: {advisor_name}")
        if entity_name: footer_parts.append(f"Entity: {entity_name}")
        if ia_reg_no: footer_parts.append(f"Reg No: {ia_reg_no}")
        footer_text = " , ".join(footer_parts)
        
        # Footer text on left
        canvas.drawString(0.5 * inch, 0.4 * inch, footer_text)
        
        # Page Number on right
        canvas.setFont("Helvetica", 8)
        page_num = canvas.getPageNumber()
        page_text = f"Page {page_num}"
        canvas.drawRightString(letter[0] - 0.5 * inch, 0.4 * inch, page_text)
        canvas.restoreState()

    @staticmethod
    def generate_blank_pdf(ia_master: Optional[IAMaster], ia_logo_path: Optional[str] = None) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.8*inch)
        
        # Attach footer data to doc
        doc.advisor_name = getattr(ia_master, 'name_of_ia', None) if ia_master else None
        doc.entity_name = getattr(ia_master, 'name_of_entity', None) if ia_master else None
        doc.ia_reg_no = getattr(ia_master, 'ia_registration_number', None) if ia_master else None
        story = []
        styles = getSampleStyleSheet()

        # Custom styles (Same as generate_pdf)
        cover_title_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#1a2980'),
            alignment=1,
            spaceAfter=40,
            fontName="Helvetica-Bold",
            leading=34
        )
        cover_subtitle_style = ParagraphStyle('CoverSubTitle', parent=styles['Normal'], fontSize=16, textColor=colors.HexColor('#45B7D1'), alignment=1, spaceAfter=60, fontName="Helvetica")
        cover_client_style = ParagraphStyle('CoverClient', parent=styles['Normal'], fontSize=18, textColor=colors.black, alignment=1, spaceAfter=15, fontName="Helvetica-Bold")
        cover_info_style = ParagraphStyle('CoverInfo', parent=styles['Normal'], fontSize=12, textColor=colors.grey, alignment=1, spaceAfter=10, fontName="Helvetica")
        heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#1a2980'), spaceAfter=12, spaceBefore=20)
        subheading_style = ParagraphStyle('SubheadingStyle', parent=styles['Heading3'], fontSize=12, textColor=colors.HexColor('#2a5298'), spaceAfter=8, spaceBefore=15)
        normal_style = ParagraphStyle('NormalStyle', parent=styles['Normal'], fontSize=10, spaceAfter=6)

        # --- COVER PAGE ---
        story.append(Spacer(1, 1.5*inch))
        if ia_logo_path and os.path.exists(ia_logo_path):
            try:
                story.append(Image(ia_logo_path, width=2.5*inch, height=1.25*inch))
            except: pass
        
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("ASSET ALLOCATION FORM", cover_title_style))
        story.append(Paragraph("Strategic Portfolio Distribution Template", cover_subtitle_style))
        story.append(Spacer(1, 0.5*inch))
        
        story.append(Paragraph("Client Name: _________________________________", cover_client_style))
        story.append(Paragraph("Client Code: _________________________________", cover_info_style))
        story.append(Paragraph("Target Risk Profile: __________________________", cover_info_style))
        
        story.append(Spacer(1, 1.2*inch))
        if ia_master:
            story.append(Paragraph(f"<b>Investment Advisor:</b>", cover_info_style))
            story.append(Paragraph(f"{ia_master.name_of_ia}", cover_info_style))
            story.append(Paragraph(f"Registration No: {ia_master.ia_registration_number}", cover_info_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Date: ________________________", cover_info_style))
        story.append(PageBreak())

        # --- BLANK TABLES ---
        story.append(Paragraph("ASSET ALLOCATION DETAILS", heading_style))
        
        story.append(Paragraph("MAIN ASSET CLASS ALLOCATION", subheading_style))
        main_data = [
            ["Asset Class", "Allocation %"],
            ["Equities", "__________%"],
            ["Debt Securities", "__________%"],
            ["Commodities", "__________%"],
            ["TOTAL", "100.0%"]
        ]
        main_table = Table(main_data, colWidths=[2.5*inch, 1.5*inch])
        main_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(main_table)

        story.append(Paragraph("SUB-ASSET CLASSES (EQUITIES)", subheading_style))
        eq_data = [
            ["Sub-asset", "Allocation %", "Within Equities", "Within Total Portfolio"],
            ["Stocks", "______%", "______%", "______%"],
            ["Mutual Funds (Equity)", "______%", "______%", "______%"],
            ["ULIP (Equity)", "______%", "______%", "______%"],
            ["TOTAL", "100.0%", "100.0%", "______%"]
        ]
        eq_table = Table(eq_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
        eq_table.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('FONTSIZE', (0,0), (-1,-1), 8)]))
        story.append(eq_table)

        story.append(Paragraph("SUB-ASSET CLASSES (DEBT)", subheading_style))
        debt_data = [
            ["Sub-asset", "Allocation %", "Within Debt", "Within Total Portfolio"],
            ["Fixed Deposits & Bonds", "______%", "______%", "______%"],
            ["Mutual Funds (Debt)", "______%", "______%", "______%"],
            ["ULIP (Debt)", "______%", "______%", "______%"],
            ["TOTAL", "100.0%", "100.0%", "______%"]
        ]
        debt_table = Table(debt_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
        debt_table.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('FONTSIZE', (0,0), (-1,-1), 8)]))
        story.append(debt_table)

        story.append(Paragraph("SUB-ASSET CLASSES (COMMODITIES)", subheading_style))
        comm_data = [
            ["Sub-asset", "Allocation %", "Within Commodities", "Within Total Portfolio"],
            ["Gold ETF", "______%", "______%", "______%"],
            ["Silver ETF", "______%", "______%", "______%"],
            ["TOTAL", "100.0%", "100.0%", "______%"]
        ]
        comm_table = Table(comm_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
        comm_table.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey), ('FONTSIZE', (0,0), (-1,-1), 8)]))
        story.append(comm_table)

        # Blank Space for Conclusion/Notes
        story.append(Paragraph("Discussions", heading_style))
        story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 2.0*inch))

        story.append(Paragraph("DISCLAIMER", heading_style))
        story.append(Paragraph(DEFAULT_ASSET_ALLOCATION_DISCLAIMER, styles['Italic']))

        # Signatures
        story.append(Spacer(1, 0.5*inch))
        sig_data = [
            [Paragraph("__________________________<br/><br/><b>Client Signature</b><br/><br/>Date: ________________", normal_style),
             Paragraph("__________________________<br/><br/><b>Advisor Signature</b><br/><br/>Date: ________________", normal_style)],
            [Paragraph("__________________________", normal_style),
             Paragraph(f"{ia_master.name_of_ia if ia_master else 'Investment Advisor'}", normal_style)]
        ]
        sig_table = Table(sig_data, colWidths=[3.5*inch, 3.5*inch])
        sig_table.setStyle(TableStyle([('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'BOTTOM'), ('BOTTOMPADDING', (0,0), (-1,0), 30)]))
        story.append(sig_table)

        doc.build(story, onFirstPage=AssetAllocationReportUtils.add_page_number, onLaterPages=AssetAllocationReportUtils.add_page_number)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_pdf(allocation: AssetAllocation, ia_master: Optional[IAMaster], ia_logo_path: Optional[str] = None) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.8*inch)
        
        # Attach footer data to doc
        doc.advisor_name = getattr(ia_master, 'name_of_ia', None) if ia_master else None
        doc.entity_name = getattr(ia_master, 'name_of_entity', None) if ia_master else None
        doc.ia_reg_no = getattr(ia_master, 'ia_registration_number', None) if ia_master else None
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        # Custom styles
        cover_title_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#1a2980'),
            alignment=1,
            spaceAfter=40,
            fontName="Helvetica-Bold",
            leading=34
        )

        cover_subtitle_style = ParagraphStyle(
            'CoverSubTitle',
            parent=styles['Normal'],
            fontSize=16,
            textColor=colors.HexColor('#45B7D1'),
            alignment=1,
            spaceAfter=60,
            fontName="Helvetica"
        )

        cover_client_style = ParagraphStyle(
            'CoverClient',
            parent=styles['Normal'],
            fontSize=18,
            textColor=colors.black,
            alignment=1,
            spaceAfter=15,
            fontName="Helvetica-Bold"
        )

        cover_info_style = ParagraphStyle(
            'CoverInfo',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.grey,
            alignment=1,
            spaceAfter=10,
            fontName="Helvetica"
        )

        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1a2980'),
            spaceAfter=12,
            spaceBefore=20
        )

        subheading_style = ParagraphStyle(
            'SubheadingStyle',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#2a5298'),
            spaceAfter=8,
            spaceBefore=15
        )

        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )

        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a2980'),
            spaceAfter=20,
            alignment=1
        )

        # --- COVER PAGE ---
        story.append(Spacer(1, 1.5*inch))
        
        # Centered Logo on Cover
        if ia_logo_path and os.path.exists(ia_logo_path):
            try:
                logo_img = Image(ia_logo_path, width=2.5*inch, height=1.25*inch)
                story.append(logo_img)
            except:
                pass
        
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("ASSET ALLOCATION REPORT", cover_title_style))
        story.append(Paragraph("Strategic Portfolio Distribution Details", cover_subtitle_style))
        
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"Client: {allocation.client.client_name}", cover_client_style))
        story.append(Paragraph(f"Code: {allocation.client.client_code.upper()}", cover_info_style))
        story.append(Paragraph(f"Risk Profile: {allocation.assigned_risk_tier}", cover_info_style))
        
        story.append(Spacer(1, 1.2*inch))
        
        if ia_master:
            story.append(Paragraph(f"<b>Prepared By:</b>", cover_info_style))
            story.append(Paragraph(f"{ia_master.name_of_ia}", cover_info_style))
            story.append(Paragraph(f"SEBI Registration No: {ia_master.ia_registration_number}", cover_info_style))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Report Date: {datetime.now().strftime('%d %B, %Y')}", cover_info_style))
        
        story.append(PageBreak())

        # --- MAIN REPORT CONTENT ---
        # The report starts directly on page 2 after the cover page.
        story.append(Paragraph("ASSET ALLOCATION SUMMARY", heading_style))
        
        # Main Allocation
        story.append(Paragraph("MAIN ASSET CLASS ALLOCATION", subheading_style))
        main_data = [
            ["Asset Class", "Allocation %"],
            ["Equities", f"{allocation.equities_percentage:.1f}%"],
            ["Debt Securities", f"{allocation.debt_securities_percentage:.1f}%"],
            ["Commodities", f"{allocation.commodities_percentage:.1f}%"],
            ["TOTAL", f"{allocation.total_allocation:.1f}%"]
        ]
        main_table = Table(main_data, colWidths=[2.5*inch, 1.5*inch])
        main_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('BACKGROUND', (0,-1), (-1,-1), colors.whitesmoke),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(main_table)
        
        # Pie Chart for Main Allocation
        if allocation.total_allocation > 0:
            labels = []
            sizes = []
            colors_list = []
            if allocation.equities_percentage > 0:
                labels.append('Equities')
                sizes.append(allocation.equities_percentage)
                colors_list.append('#FF6B6B')
            if allocation.debt_securities_percentage > 0:
                labels.append('Debt')
                sizes.append(allocation.debt_securities_percentage)
                colors_list.append('#4ECDC4')
            if allocation.commodities_percentage > 0:
                labels.append('Commodities')
                sizes.append(allocation.commodities_percentage)
                colors_list.append('#45B7D1')
            
            if sizes:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(labels, sizes, "Main Asset Class Allocation", colors_list)
                img = Image(io.BytesIO(chart_bytes), width=2.5*inch, height=2*inch)
                story.append(Spacer(1, 10))
                story.append(img)

        # Sub-Asset Details (Equities)
        if allocation.equities_percentage > 0:
            story.append(Paragraph("EQUITIES SUB-ASSET ALLOCATION", subheading_style))
            eq_data = [
                ["Sub-asset", "Allocation %", "Within Equities", "Within Total Portfolio"],
                ["Stocks", f"{allocation.stocks_percentage:.1f}%", f"{allocation.stocks_percentage:.1f}%", f"{(allocation.stocks_percentage * allocation.equities_percentage / 100):.1f}%"],
                ["Mutual Funds (Equity)", f"{allocation.mutual_fund_equity_percentage:.1f}%", f"{allocation.mutual_fund_equity_percentage:.1f}%", f"{(allocation.mutual_fund_equity_percentage * allocation.equities_percentage / 100):.1f}%"],
                ["ULIP (Equity)", f"{allocation.ulip_equity_percentage:.1f}%", f"{allocation.ulip_equity_percentage:.1f}%", f"{(allocation.ulip_equity_percentage * allocation.equities_percentage / 100):.1f}%"],
                ["TOTAL", "100.0%", "100.0%", f"{allocation.equities_percentage:.1f}%"]
            ]
            eq_table = Table(eq_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
            eq_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('PADDING', (0,0), (-1,-1), 4),
                ('FONTSIZE', (0,0), (-1,-1), 8),
            ]))
            story.append(eq_table)

            # Pie Chart for Equities
            eq_labels = []
            eq_sizes = []
            if allocation.stocks_percentage > 0: eq_labels.append('Stocks'); eq_sizes.append(allocation.stocks_percentage)
            if allocation.mutual_fund_equity_percentage > 0: eq_labels.append('Mutual Funds'); eq_sizes.append(allocation.mutual_fund_equity_percentage)
            if allocation.ulip_equity_percentage > 0: eq_labels.append('ULIP'); eq_sizes.append(allocation.ulip_equity_percentage)
            
            if eq_sizes:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(eq_labels, eq_sizes, "Equities Sub-Asset Allocation", ['#ef4444', '#f06565', '#f38787'])
                img = Image(io.BytesIO(chart_bytes), width=2.5*inch, height=2*inch)
                story.append(Spacer(1, 10))
                story.append(img)
            story.append(Spacer(1, 15))

        # Sub-Asset Details (Debt)
        if allocation.debt_securities_percentage > 0:
            story.append(Paragraph("DEBT SECURITIES SUB-ASSET ALLOCATION", subheading_style))
            debt_data = [
                ["Sub-asset", "Allocation %", "Within Debt", "Within Total Portfolio"],
                ["Fixed Deposits & Bonds", f"{allocation.fixed_deposits_bonds_percentage:.1f}%", f"{allocation.fixed_deposits_bonds_percentage:.1f}%", f"{(allocation.fixed_deposits_bonds_percentage * allocation.debt_securities_percentage / 100):.1f}%"],
                ["Mutual Funds (Debt)", f"{allocation.mutual_fund_debt_percentage:.1f}%", f"{allocation.mutual_fund_debt_percentage:.1f}%", f"{(allocation.mutual_fund_debt_percentage * allocation.debt_securities_percentage / 100):.1f}%"],
                ["ULIP (Debt)", f"{allocation.ulip_debt_percentage:.1f}%", f"{allocation.ulip_debt_percentage:.1f}%", f"{(allocation.ulip_debt_percentage * allocation.debt_securities_percentage / 100):.1f}%"],
                ["TOTAL", "100.0%", "100.0%", f"{allocation.debt_securities_percentage:.1f}%"]
            ]
            debt_table = Table(debt_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
            debt_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('PADDING', (0,0), (-1,-1), 4),
                ('FONTSIZE', (0,0), (-1,-1), 8),
            ]))
            story.append(debt_table)

            # Pie Chart for Debt
            debt_labels = []
            debt_sizes = []
            if allocation.fixed_deposits_bonds_percentage > 0: debt_labels.append('FD/Bonds'); debt_sizes.append(allocation.fixed_deposits_bonds_percentage)
            if allocation.mutual_fund_debt_percentage > 0: debt_labels.append('Mutual Funds'); debt_sizes.append(allocation.mutual_fund_debt_percentage)
            if allocation.ulip_debt_percentage > 0: debt_labels.append('ULIP'); debt_sizes.append(allocation.ulip_debt_percentage)
            
            if debt_sizes:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(debt_labels, debt_sizes, "Debt Sub-Asset Allocation", ['#3b82f6', '#619bf8', '#88b4fa'])
                img = Image(io.BytesIO(chart_bytes), width=2.5*inch, height=2*inch)
                story.append(Spacer(1, 10))
                story.append(img)
            story.append(Spacer(1, 15))

        # Sub-Asset Details (Commodities)
        if allocation.commodities_percentage > 0:
            story.append(Paragraph("COMMODITIES SUB-ASSET ALLOCATION", subheading_style))
            comm_data = [
                ["Sub-asset", "Allocation %", "Within Commodities", "Within Total Portfolio"],
                ["Gold ETF", f"{allocation.gold_etf_percentage:.1f}%", f"{allocation.gold_etf_percentage:.1f}%", f"{(allocation.gold_etf_percentage * allocation.commodities_percentage / 100):.1f}%"],
                ["Silver ETF", f"{allocation.silver_etf_percentage:.1f}%", f"{allocation.silver_etf_percentage:.1f}%", f"{(allocation.silver_etf_percentage * allocation.commodities_percentage / 100):.1f}%"],
                ["TOTAL", "100.0%", "100.0%", f"{allocation.commodities_percentage:.1f}%"]
            ]
            comm_table = Table(comm_data, colWidths=[2.3*inch, 1.2*inch, 1.2*inch, 1.8*inch])
            comm_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('PADDING', (0,0), (-1,-1), 4),
                ('FONTSIZE', (0,0), (-1,-1), 8),
            ]))
            story.append(comm_table)

            # Pie Chart for Commodities
            comm_labels = []
            comm_sizes = []
            if allocation.gold_etf_percentage > 0: comm_labels.append('Gold ETF'); comm_sizes.append(allocation.gold_etf_percentage)
            if allocation.silver_etf_percentage > 0: comm_labels.append('Silver ETF'); comm_sizes.append(allocation.silver_etf_percentage)
            
            if comm_sizes:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(comm_labels, comm_sizes, "Commodities Sub-Asset Allocation", ['#f59e0b', '#f7b13c', '#fac56d'])
                img = Image(io.BytesIO(chart_bytes), width=2.5*inch, height=2*inch)
                story.append(Spacer(1, 10))
                story.append(img)
            story.append(Spacer(1, 15))

        # Advisor Recommendation
        if allocation.tier_recommendation:
            story.append(Paragraph("ADVISOR RECOMMENDATION", heading_style))
            story.append(Paragraph(allocation.tier_recommendation, normal_style))
            story.append(Spacer(1, 10))

        # System Conclusion
        if allocation.system_conclusion:
            story.append(Paragraph("CONCLUSION", heading_style))
            for part in allocation.system_conclusion.split('\n\n'):
                story.append(Paragraph(part.replace('\n', '<br/>'), normal_style))
                story.append(Spacer(1, 6))

        # Disclaimer
        story.append(Paragraph("DISCLAIMER", heading_style))
        story.append(Paragraph(DEFAULT_ASSET_ALLOCATION_DISCLAIMER, styles['Italic']))
        
        if allocation.disclaimer_text:
            story.append(Spacer(1, 10))
            story.append(Paragraph(allocation.disclaimer_text, styles['Italic']))

        # Discussion Notes
        if allocation.discussion_notes:
            story.append(Paragraph("DISCUSSION NOTES", heading_style))
            story.append(Paragraph(allocation.discussion_notes, normal_style))
            story.append(Spacer(1, 15))

        # --- SIGNATURE SECTION ---
        story.append(Spacer(1, 0.5*inch))
        # Style for signatures with more spacing
        sig_style = ParagraphStyle(
            'SigStyle',
            parent=normal_style,
            leading=16  # Increased vertical spacing
        )
        
        sig_data = [
            [
                Paragraph("__________________________<br/><br/><b>Client Signature</b><br/><br/>Date: ________________", sig_style),
                Paragraph("__________________________<br/><br/><b>Advisor Signature</b><br/><br/>Date: ________________", sig_style)
            ],
            [
                Paragraph(f"{allocation.client.client_name}", sig_style),
                Paragraph(f"{ia_master.name_of_ia if ia_master else 'Investment Advisor'}", sig_style)
            ]
        ]
        sig_table = Table(sig_data, colWidths=[3.5*inch, 3.5*inch])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
            ('BOTTOMPADDING', (0,0), (-1,0), 30),
        ]))
        story.append(sig_table)

        doc.build(story, onFirstPage=AssetAllocationReportUtils.add_page_number, onLaterPages=AssetAllocationReportUtils.add_page_number)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_docx(allocation: AssetAllocation, ia_master: Optional[IAMaster]) -> io.BytesIO:
        doc = Document()
        
        # Word Header (Simplified)
        section = doc.sections[0]
        header = section.header
        htable = header.add_table(1, 2, doc.sections[0].page_width)
        
        if ia_master:
            htable.cell(0, 1).text = f"{ia_master.name_of_ia}\nReg: {ia_master.ia_registration_number}\n{ia_master.registered_email_id}"
            htable.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_heading('ASSET ALLOCATION REPORT', 0)
        
        # 1. Client Table
        doc.add_heading('Client Details', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        table.cell(0, 0).text = "Client Name"
        table.cell(0, 1).text = allocation.client.client_name
        table.cell(1, 0).text = "Client Code"
        table.cell(1, 1).text = allocation.client.client_code.upper()
        table.cell(2, 0).text = "Risk Tier"
        table.cell(2, 1).text = allocation.assigned_risk_tier
        table.cell(3, 0).text = "Date"
        table.cell(3, 1).text = allocation.created_at.strftime('%d %B, %Y')
        
        doc.add_paragraph()
        
        # 2. Main Allocation Table & Chart
        doc.add_heading('MAIN ASSET CLASS ALLOCATION', level=1)
        atable = doc.add_table(rows=4, cols=2)
        atable.style = 'Medium List 1 Accent 1'
        atable.cell(0, 0).text = "Asset Class"
        atable.cell(0, 1).text = "Percentage"
        
        data = [
            ("Equities", f"{allocation.equities_percentage:.1f}%"),
            ("Debt", f"{allocation.debt_securities_percentage:.1f}%"),
            ("Commodities", f"{allocation.commodities_percentage:.1f}%")
        ]
        for i, (label, val) in enumerate(data, 1):
            atable.cell(i, 0).text = label
            atable.cell(i, 1).text = val

        # Chart for Main Allocation
        if allocation.total_allocation > 0:
            labels, sizes, colors_list = [], [], []
            if allocation.equities_percentage > 0: 
                labels.append('Equities'); sizes.append(allocation.equities_percentage); colors_list.append('#FF6B6B')
            if allocation.debt_securities_percentage > 0: 
                labels.append('Debt'); sizes.append(allocation.debt_securities_percentage); colors_list.append('#4ECDC4')
            if allocation.commodities_percentage > 0: 
                labels.append('Commodities'); sizes.append(allocation.commodities_percentage); colors_list.append('#45B7D1')
            
            if sizes:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(labels, sizes, "Main Asset Class Allocation", colors_list)
                doc.add_picture(io.BytesIO(chart_bytes), width=Inches(4.5))
                doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3. Equities Breakdown
        if allocation.equities_percentage > 0:
            doc.add_heading('EQUITIES SUB-ASSET ALLOCATION', level=2)
            eq_table = doc.add_table(rows=5, cols=4)
            eq_table.style = 'Table Grid'
            hdr = eq_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Sub-asset", "Allocation %", "Within Equities", "Within Total Portfolio"
            
            eq_data = [
                ("Stocks", f"{allocation.stocks_percentage:.1f}%", f"{allocation.stocks_percentage:.1f}%", f"{(allocation.stocks_percentage * allocation.equities_percentage / 100):.1f}%"),
                ("Mutual Funds", f"{allocation.mutual_fund_equity_percentage:.1f}%", f"{allocation.mutual_fund_equity_percentage:.1f}%", f"{(allocation.mutual_fund_equity_percentage * allocation.equities_percentage / 100):.1f}%"),
                ("ULIP", f"{allocation.ulip_equity_percentage:.1f}%", f"{allocation.ulip_equity_percentage:.1f}%", f"{(allocation.ulip_equity_percentage * allocation.equities_percentage / 100):.1f}%"),
                ("TOTAL", "100.0%", "100.0%", f"{allocation.equities_percentage:.1f}%")
            ]
            for i, (s, a, we, wp) in enumerate(eq_data, 1):
                row = eq_table.rows[i].cells
                row[0].text, row[1].text, row[2].text, row[3].text = s, a, we, wp

            # Chart for Equities
            eq_l, eq_s = [], []
            if allocation.stocks_percentage > 0: eq_l.append('Stocks'); eq_s.append(allocation.stocks_percentage)
            if allocation.mutual_fund_equity_percentage > 0: eq_l.append('Mutual Funds'); eq_s.append(allocation.mutual_fund_equity_percentage)
            if allocation.ulip_equity_percentage > 0: eq_l.append('ULIP'); eq_s.append(allocation.ulip_equity_percentage)
            if eq_s:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(eq_l, eq_s, "Equities Breakdown", ['#ef4444', '#f06565', '#f38787'])
                doc.add_picture(io.BytesIO(chart_bytes), width=Inches(3.5))

        # 4. Debt Breakdown
        if allocation.debt_securities_percentage > 0:
            doc.add_heading('DEBT SECURITIES SUB-ASSET ALLOCATION', level=2)
            db_table = doc.add_table(rows=5, cols=4)
            db_table.style = 'Table Grid'
            hdr = db_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Sub-asset", "Allocation %", "Within Debt", "Within Total Portfolio"
            
            db_data = [
                ("FD/Bonds", f"{allocation.fixed_deposits_bonds_percentage:.1f}%", f"{allocation.fixed_deposits_bonds_percentage:.1f}%", f"{(allocation.fixed_deposits_bonds_percentage * allocation.debt_securities_percentage / 100):.1f}%"),
                ("Mutual Funds", f"{allocation.mutual_fund_debt_percentage:.1f}%", f"{allocation.mutual_fund_debt_percentage:.1f}%", f"{(allocation.mutual_fund_debt_percentage * allocation.debt_securities_percentage / 100):.1f}%"),
                ("ULIP", f"{allocation.ulip_debt_percentage:.1f}%", f"{allocation.ulip_debt_percentage:.1f}%", f"{(allocation.ulip_debt_percentage * allocation.debt_securities_percentage / 100):.1f}%"),
                ("TOTAL", "100.0%", "100.0%", f"{allocation.debt_securities_percentage:.1f}%")
            ]
            for i, (s, a, wd, wp) in enumerate(db_data, 1):
                row = db_table.rows[i].cells
                row[0].text, row[1].text, row[2].text, row[3].text = s, a, wd, wp

            # Chart for Debt
            db_l, db_s = [], []
            if allocation.fixed_deposits_bonds_percentage > 0: db_l.append('FD/Bonds'); db_s.append(allocation.fixed_deposits_bonds_percentage)
            if allocation.mutual_fund_debt_percentage > 0: db_l.append('Mutual Funds'); db_s.append(allocation.mutual_fund_debt_percentage)
            if allocation.ulip_debt_percentage > 0: db_l.append('ULIP'); db_s.append(allocation.ulip_debt_percentage)
            if db_s:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(db_l, db_s, "Debt Breakdown", ['#3b82f6', '#619bf8', '#88b4fa'])
                doc.add_picture(io.BytesIO(chart_bytes), width=Inches(3.5))

        # 5. Commodities Breakdown
        if allocation.commodities_percentage > 0:
            doc.add_heading('COMMODITIES SUB-ASSET ALLOCATION', level=2)
            cm_table = doc.add_table(rows=4, cols=4)
            cm_table.style = 'Table Grid'
            hdr = cm_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Sub-asset", "Allocation %", "Within Commodities", "Within Total Portfolio"
            
            cm_data = [
                ("Gold ETF", f"{allocation.gold_etf_percentage:.1f}%", f"{allocation.gold_etf_percentage:.1f}%", f"{(allocation.gold_etf_percentage * allocation.commodities_percentage / 100):.1f}%"),
                ("Silver ETF", f"{allocation.silver_etf_percentage:.1f}%", f"{allocation.silver_etf_percentage:.1f}%", f"{(allocation.silver_etf_percentage * allocation.commodities_percentage / 100):.1f}%"),
                ("TOTAL", "100.0%", "100.0%", f"{allocation.commodities_percentage:.1f}%")
            ]
            for i, (s, a, wc, wp) in enumerate(cm_data, 1):
                row = cm_table.rows[i].cells
                row[0].text, row[1].text, row[2].text, row[3].text = s, a, wc, wp

            # Chart for Commodities
            cm_l, cm_s = [], []
            if allocation.gold_etf_percentage > 0: cm_l.append('Gold'); cm_s.append(allocation.gold_etf_percentage)
            if allocation.silver_etf_percentage > 0: cm_l.append('Silver'); cm_s.append(allocation.silver_etf_percentage)
            if cm_s:
                chart_bytes = AssetAllocationReportUtils.create_pie_chart(cm_l, cm_s, "Commodities Breakdown", ['#f59e0b', '#f7b13c', '#fac56d'])
                doc.add_picture(io.BytesIO(chart_bytes), width=Inches(3.5))

        # 6. Advisor Recommendation & Conclusion
        if allocation.tier_recommendation:
            doc.add_heading('ADVISOR RECOMMENDATION', level=1)
            doc.add_paragraph(allocation.tier_recommendation)

        if allocation.system_conclusion:
            doc.add_heading('CONCLUSION', level=1)
            doc.add_paragraph(allocation.system_conclusion)

        if allocation.discussion_notes:
            doc.add_heading('DISCUSSION NOTES', level=1)
            doc.add_paragraph(allocation.discussion_notes)

        doc.add_heading('DISCLAIMER', level=2)
        doc.add_paragraph(DEFAULT_ASSET_ALLOCATION_DISCLAIMER).italic = True
        
        if allocation.disclaimer_text:
            doc.add_paragraph(allocation.disclaimer_text).italic = True

        # 7. Signatures for Word
        doc.add_paragraph('\n\n')
        stable = doc.add_table(rows=2, cols=2)
        stable.cell(0, 0).text = "__________________________\n\nClient Signature\n\nDate: ________________"
        stable.cell(0, 1).text = "__________________________\n\nAdvisor Signature\n\nDate: ________________"
        stable.cell(1, 0).text = allocation.client.client_name
        stable.cell(1, 1).text = ia_master.name_of_ia if ia_master else "Investment Advisor"

        # Add Page Footer
        section = doc.sections[0]
        footer = section.footer
        f_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        prepared_by = getattr(ia_master, 'name_of_ia', 'INVESTMENT ADVISOR') if ia_master else 'INVESTMENT ADVISOR'
        ia_entity = getattr(ia_master, 'name_of_entity', 'N/A') if ia_master else 'N/A'
        ia_reg = getattr(ia_master, 'ia_registration_number', 'N/A') if ia_master else 'N/A'
        
        footer_text = f"Prepared by: {prepared_by} , Entity: {ia_entity} , Reg No: {ia_reg}"
        f_run = f_p.add_run(footer_text)
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        f_run.italic = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

